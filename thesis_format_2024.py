import copy
import os
import re
import sys
import tempfile
import zipfile

from thesis_config import DEFAULT_CONFIG, resolve_config, resolve_logo_path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_JC_MAP = {
    "left": "left",
    "center": "center",
    "right": "right",
    "justify": "both",
}


def _resource_dir():
    """PyInstaller bundle or script directory."""
    return getattr(sys, "_MEIPASS", os.path.dirname(os.path.abspath(__file__)))


def set_rfonts(rpr, east_asia, latin="Times New Roman"):
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def set_style_font(style, east_asia, size_pt, bold=False, latin="Times New Roman"):
    style.font.name = latin
    style.font.size = size_pt
    if bold is not None:
        style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)
    rpr = style.element.get_or_add_rPr()
    set_rfonts(rpr, east_asia, latin)


def set_run_font(run, east_asia, size_pt=None, bold=None, latin="Times New Roman"):
    run.font.name = latin
    if size_pt is not None:
        run.font.size = size_pt
    if bold is not None:
        run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)
    rpr = run._element.get_or_add_rPr()
    set_rfonts(rpr, east_asia, latin)


def set_para_runs_font(para, east_asia, size_pt, bold=None, latin="Times New Roman"):
    for run in para.runs:
        set_run_font(run, east_asia=east_asia, size_pt=size_pt, bold=bold, latin=latin)


def zero_spacing(para):
    pf = para.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def set_table_border(cell, edge, sz, val="single", color="000000"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    edge_el = tc_borders.find(qn(f"w:{edge}"))
    if edge_el is None:
        edge_el = OxmlElement(f"w:{edge}")
        tc_borders.append(edge_el)
    edge_el.set(qn("w:val"), val)
    edge_el.set(qn("w:sz"), str(sz))
    edge_el.set(qn("w:color"), color)


def clear_table_border(cell, edge):
    set_table_border(cell, edge, sz=0, val="nil")


def _ensure_keep_next(p_el):
    """Set keepNext on a w:p element to prevent page break after it."""
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    if pPr.find(qn("w:keepNext")) is None:
        pPr.append(OxmlElement("w:keepNext"))


def _set_para_spacing(p_el, side, pt_val):
    """Set w:spacing before/after on a w:p XML element. pt_val is a Pt() value."""
    pPr = p_el.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    twips = str(int(pt_val.pt * 20))
    spacing.set(qn("w:" + side), twips)


def _check_caption_numbering(doc, fig_pat, tbl_pat, cfg=None):
    """Validate sequential numbering of figure/table captions (body + appendix)."""
    warnings = []
    sec = cfg.get("sections", {}) if cfg else {}
    appendix_re = re.compile(sec.get("appendix_pattern", r"^附录\s*[A-Z]"))

    current_appendix = None  # None=body, "A"/"B"=appendix letter
    body_figs, body_tbls = [], []
    app_figs = {}  # {"A": [1,2], "B": [1]}
    app_tbls = {}

    for para in doc.paragraphs:
        t = para.text.strip()
        sn = para.style.name if para.style else ""

        # Detect appendix heading → switch to appendix mode
        if sn in ("Heading 1", "样式1") and appendix_re.match(t):
            m = re.search(r"附录\s*([A-Z])", t)
            if m:
                current_appendix = m.group(1)
            continue

        if current_appendix:
            # Appendix-style captions: 图A1, 表B2, etc.
            m = re.match(r"^图([A-Z])(\d+)", t)
            if m:
                app_figs.setdefault(m.group(1), []).append(int(m.group(2)))
                continue
            m = re.match(r"^(续)?表([A-Z])(\d+)", t)
            if m:
                if m.group(1):  # 续表 skip
                    continue
                app_tbls.setdefault(m.group(2), []).append(int(m.group(3)))
                continue
        else:
            # Body captions
            if re.match(fig_pat, t) or re.match(r"^Figure\s*\d", t, re.I):
                m = re.search(r"(\d+)", t)
                if m:
                    body_figs.append(int(m.group(1)))
            elif re.match(tbl_pat, t) or re.match(r"^Table\s*\d", t, re.I):
                if re.match(r"^续", t):
                    continue
                m = re.search(r"(\d+)", t)
                if m:
                    body_tbls.append(int(m.group(1)))

    # Validate body
    for label, nums in [("图", body_figs), ("表", body_tbls)]:
        if not nums:
            continue
        expected = list(range(1, len(nums) + 1))
        if nums != expected:
            warnings.append(f"  警告: 正文{label}编号不连续 — 期望 {expected}, 实际 {nums}")
    # Validate appendix
    for letter in sorted(set(list(app_figs.keys()) + list(app_tbls.keys()))):
        for label, store in [("图", app_figs), ("表", app_tbls)]:
            nums = store.get(letter, [])
            if not nums:
                continue
            expected = list(range(1, len(nums) + 1))
            if nums != expected:
                warnings.append(
                    f"  警告: 附录{letter}{label}编号不连续 — 期望 {expected}, 实际 {nums}")
    return warnings


def is_heading(para, level):
    return para.style and para.style.name == f"Heading {level}"


def contains_cjk(s):
    return any("\u4e00" <= ch <= "\u9fff" for ch in s)


def normalize_cn_keywords(text):
    m = re.match(r"^\s*关键词\s*[：:]\s*(.*)$", text)
    if not m:
        return None
    items = [x.strip(" ；;。,. ") for x in re.split(r"[；;]", m.group(1)) if x.strip(" ；;。,. ")]
    return "关键词：" + "；".join(items)


def cap_token(token):
    if "-" in token:
        parts = token.split("-")
        out = []
        for p in parts:
            if re.search(r"[A-Za-z]", p):
                out.append(p[:1].upper() + p[1:].lower())
            else:
                out.append(p)
        return "-".join(out)
    if re.search(r"[A-Za-z]", token):
        return token[:1].upper() + token[1:].lower()
    return token


def title_case_phrase(s):
    words = [w for w in re.split(r"\s+", s.strip()) if w]
    return " ".join([cap_token(w) for w in words])


def normalize_en_keywords(text):
    m = re.match(r"^\s*Key\s*words\s*:\s*(.*)$", text, flags=re.I)
    if not m:
        return None
    items = [x.strip(" ;；。,. ") for x in re.split(r"[;；]", m.group(1)) if x.strip(" ;；。,. ")]
    items = [title_case_phrase(x) for x in items]
    return "Key words: " + "; ".join(items)


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------

def normalize_sections(doc, cfg):
    pg = cfg["page"]
    for section in doc.sections:
        section.top_margin = Cm(pg["margins"]["top"])
        section.bottom_margin = Cm(pg["margins"]["bottom"])
        section.left_margin = Cm(pg["margins"]["left"])
        section.right_margin = Cm(pg["margins"]["right"])
        section.gutter = Cm(pg["gutter"])
        section.header_distance = Cm(pg["header_distance"])
        section.footer_distance = Cm(pg["footer_distance"])
        try:
            for p in section.header.paragraphs:
                p.clear()
        except Exception:
            pass


# ---------------------------------------------------------------------------
# Page numbers
# ---------------------------------------------------------------------------

def add_page_number_field(paragraph, cfg, align="center"):
    paragraph.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    pn_size = Pt(cfg["sizes"]["page_number"])
    pn_cfg = cfg["page_numbers"]
    pn_font = pn_cfg.get("font", "") or cfg["fonts"]["latin"]
    pn_bold = pn_cfg.get("bold", False)
    body_ea = cfg["fonts"]["body"]
    decorator = pn_cfg.get("decorator", "{page}")
    prefix, suffix = "", ""
    if "{page}" in decorator:
        parts = decorator.split("{page}", 1)
        prefix, suffix = parts[0], parts[1]
    elif decorator != "{page}":
        prefix = ""

    def _pn_run(text=None):
        r = paragraph.add_run(text) if text else paragraph.add_run()
        r.font.size = pn_size
        r.font.name = pn_font
        r.font.bold = pn_bold
        rpr = r._element.get_or_add_rPr()
        set_rfonts(rpr, body_ea, pn_font)
        return r

    if prefix:
        _pn_run(prefix)

    run = _pn_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._element.append(fld_char_begin)

    run2 = _pn_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run2._element.append(instr)

    run3 = _pn_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run3._element.append(fld_sep)

    _pn_run("1")

    run5 = _pn_run()
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run5._element.append(fld_char_end)

    if suffix:
        _pn_run(suffix)


def set_section_page_number_format(section, fmt="decimal", start=None):
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:fmt"), fmt)
    if start is not None:
        pg_num.set(qn("w:start"), str(start))


def insert_section_break_before(paragraph):
    p_element = paragraph._element
    prev = p_element.getprevious()
    if prev is None:
        return None
    pPr = prev.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        prev.insert(0, pPr)
    sect_pr = OxmlElement("w:sectPr")
    pPr.append(sect_pr)
    return sect_pr


def insert_page_break_after(paragraph):
    p_element = paragraph._element
    new_p = OxmlElement("w:p")
    new_r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    new_r.append(br)
    new_p.append(new_r)
    p_element.addnext(new_p)


def _enable_even_odd_headers(section):
    sect_pr = section._sectPr
    existing = sect_pr.find(qn("w:titlePg"))
    if existing is None:
        sect_pr.append(OxmlElement("w:titlePg"))


def _set_even_odd_on_doc(doc):
    settings = doc.settings.element
    if settings.find(qn("w:evenAndOddHeaders")) is None:
        settings.append(OxmlElement("w:evenAndOddHeaders"))


def _setup_single_section_pn(doc, cfg):
    """Set up page numbers when there is only one section (no front/body split)."""
    pn = cfg["page_numbers"]
    body_pos = pn.get("body_position", "center")
    body_odd = pn.get("body_odd_position", "right")
    body_even = pn.get("body_even_position", "left")
    need_alternate = body_pos == "alternate"
    hf_diff_oe = cfg.get("header_footer", {}).get("different_odd_even", False) and \
                 cfg.get("header_footer", {}).get("enabled", False)
    need_even_odd = need_alternate or hf_diff_oe

    section = doc.sections[0]
    set_section_page_number_format(section, fmt=pn["body_format"], start=pn["body_start"])

    if need_even_odd:
        _set_even_odd_on_doc(doc)

    footer = section.footer
    footer.is_linked_to_previous = False
    for p in footer.paragraphs:
        p.clear()

    if need_alternate:
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        add_page_number_field(fp, cfg, align=body_odd)
        even_footer = section.even_page_footer
        even_footer.is_linked_to_previous = False
        for p in even_footer.paragraphs:
            p.clear()
        ep = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
        add_page_number_field(ep, cfg, align=body_even)
    else:
        actual_pos = body_pos if body_pos != "alternate" else "center"
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        add_page_number_field(fp, cfg, align=actual_pos)
        if need_even_odd:
            even_footer = section.even_page_footer
            even_footer.is_linked_to_previous = False
            for p in even_footer.paragraphs:
                p.clear()
            ep = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
            add_page_number_field(ep, cfg, align=actual_pos)


def setup_page_numbers(doc, cfg):
    pn = cfg["page_numbers"]
    front_pos = pn.get("front_position", "center")
    body_pos = pn.get("body_position", "center")
    body_odd = pn.get("body_odd_position", "right")
    body_even = pn.get("body_even_position", "left")
    need_alternate = body_pos == "alternate"
    hf_diff_oe = cfg.get("header_footer", {}).get("different_odd_even", False) and \
                 cfg.get("header_footer", {}).get("enabled", False)
    need_even_odd = need_alternate or hf_diff_oe

    first_body_h1 = None
    for para in doc.paragraphs:
        if para.style and para.style.name == "Heading 1":
            t = para.text.strip()
            if t.replace(" ", "").replace("\u3000", "") == _find_special_display(cfg, "目录", raw=True):
                continue
            first_body_h1 = para
            break

    if first_body_h1 is None:
        _setup_single_section_pn(doc, cfg)
        return

    new_sect_pr = insert_section_break_before(first_body_h1)
    if new_sect_pr is None:
        _setup_single_section_pn(doc, cfg)
        return

    for attr in ["pgSz", "pgMar"]:
        existing = doc.sections[0]._sectPr.find(qn(f"w:{attr}"))
        if existing is not None:
            new_sect_pr.append(copy.deepcopy(existing))

    set_section_page_number_format(
        doc.sections[0], fmt=pn["front_format"], start=pn["front_start"])

    if len(doc.sections) > 1:
        set_section_page_number_format(
            doc.sections[-1], fmt=pn["body_format"], start=pn["body_start"])

    if need_even_odd:
        _set_even_odd_on_doc(doc)

    for idx, section in enumerate(doc.sections):
        is_body = idx == len(doc.sections) - 1 and len(doc.sections) > 1
        pos = body_pos if is_body else front_pos

        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()

        if pos == "alternate" and is_body:
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            add_page_number_field(fp, cfg, align=body_odd)
            even_footer = section.even_page_footer
            even_footer.is_linked_to_previous = False
            for p in even_footer.paragraphs:
                p.clear()
            ep = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
            add_page_number_field(ep, cfg, align=body_even)
        else:
            actual_pos = pos if pos != "alternate" else "center"
            fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            add_page_number_field(fp, cfg, align=actual_pos)
            if need_even_odd:
                even_footer = section.even_page_footer
                even_footer.is_linked_to_previous = False
                for p in even_footer.paragraphs:
                    p.clear()
                ep = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
                add_page_number_field(ep, cfg, align=actual_pos)


# ---------------------------------------------------------------------------
# Headers (页眉)
# ---------------------------------------------------------------------------

def _add_header_border(paragraph, width_pt=0.75, style="single"):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    bottom = OxmlElement("w:bottom")
    sz = int(width_pt * 8)
    val = "double" if style == "double" else "single"
    bottom.set(qn("w:val"), val)
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)


def _render_header_para(paragraph, text, cfg, hf_cfg, doc, align="center"):
    paragraph.alignment = _ALIGN_MAP.get(align, WD_ALIGN_PARAGRAPH.CENTER)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    rendered = text
    if "{chapter_title}" in rendered:
        rendered_prefix, rendered_suffix = rendered.split("{chapter_title}", 1)
        paragraph.clear()

        # Chinese Word displays "Heading 1" as "标题 1" in the UI;
        # STYLEREF matches by display name, so use the Chinese name.
        # Also add "标题 1" as alias to ensure STYLEREF can find it.
        for sty in doc.styles:
            if sty.style_id == "Heading1" or sty.name == "Heading 1":
                name_el = sty.element.find(qn("w:name"))
                if name_el is not None:
                    aliases_el = sty.element.find(qn("w:aliases"))
                    if aliases_el is None:
                        aliases_el = OxmlElement("w:aliases")
                        name_el.addnext(aliases_el)
                    existing = aliases_el.get(qn("w:val"), "")
                    if "标题 1" not in existing:
                        new_val = ("标题 1," + existing) if existing else "标题 1"
                        aliases_el.set(qn("w:val"), new_val)
                break
        h1_style_name = "标题 1"

        def _hf_run(txt=None):
            r = paragraph.add_run(txt) if txt else paragraph.add_run()
            set_run_font(r, east_asia=hf_cfg["font"],
                         size_pt=Pt(hf_cfg["font_size"]),
                         bold=hf_cfg.get("bold", False),
                         latin=cfg["fonts"]["latin"])
            return r

        if rendered_prefix:
            _hf_run(rendered_prefix)
        r1 = _hf_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        r1._element.append(begin)
        r2 = _hf_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f' STYLEREF "{h1_style_name}" '
        r2._element.append(instr)
        r3 = _hf_run()
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        r3._element.append(sep)
        _hf_run("(章标题)")
        r5 = _hf_run()
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        r5._element.append(end)
        if rendered_suffix:
            _hf_run(rendered_suffix)
    else:
        paragraph.clear()
        r = paragraph.add_run(rendered)
        set_run_font(r, east_asia=hf_cfg["font"],
                     size_pt=Pt(hf_cfg["font_size"]),
                     bold=hf_cfg.get("bold", False),
                     latin=cfg["fonts"]["latin"])
    if hf_cfg.get("border_bottom", False):
        _add_header_border(paragraph,
                           hf_cfg.get("border_bottom_width", 0.75),
                           hf_cfg.get("border_bottom_style", "single"))


def setup_headers(doc, cfg):
    hf = cfg.get("header_footer", {})
    if not hf.get("enabled", False):
        return
    scope = hf.get("scope", "body")
    diff_oe = hf.get("different_odd_even", True)
    first_no = hf.get("first_page_no_header", False)
    odd_align = hf.get("odd_page_align", "center")
    even_align = hf.get("even_page_align", "center")

    if diff_oe:
        _set_even_odd_on_doc(doc)
    doc_has_even_odd = doc.settings.element.find(qn("w:evenAndOddHeaders")) is not None

    for idx, section in enumerate(doc.sections):
        is_front = idx == 0 and len(doc.sections) > 1
        if scope == "body" and is_front:
            continue

        if first_no:
            sect_pr = section._sectPr
            if sect_pr.find(qn("w:titlePg")) is None:
                sect_pr.append(OxmlElement("w:titlePg"))

        header = section.header
        header.is_linked_to_previous = False
        for p in header.paragraphs:
            p.clear()
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        _render_header_para(hp, hf["odd_page_text"], cfg, hf, doc, align=odd_align)

        if doc_has_even_odd:
            even_header = section.even_page_header
            even_header.is_linked_to_previous = False
            for p in even_header.paragraphs:
                p.clear()
            ep = even_header.paragraphs[0] if even_header.paragraphs else even_header.add_paragraph()
            even_text = hf["even_page_text"] if diff_oe else hf["odd_page_text"]
            even_a = even_align if diff_oe else odd_align
            _render_header_para(ep, even_text, cfg, hf, doc, align=even_a)

        if first_no:
            first_header = section.first_page_header
            first_header.is_linked_to_previous = False
            for p in first_header.paragraphs:
                p.clear()
            # Preserve first-page footer (page number) by unlinking it
            first_footer = section.first_page_footer
            first_footer.is_linked_to_previous = False


# ---------------------------------------------------------------------------
# Special title helpers
# ---------------------------------------------------------------------------

def _detect_front_matter(doc, cfg):
    """Detect whether the document has front matter (abstract/keywords) before first H1."""
    sec = cfg.get("sections", {})
    cn_kw_re = sec.get("cn_keywords_pattern", r"^\s*关键词\s*[：:]")
    en_abs_re = sec.get("en_abstract_pattern", r"(?i)^\s*Abstract\s*[：:]")
    toc_display = _find_special_display(cfg, "目录", raw=True)

    first_h1_idx = None
    for i, para in enumerate(doc.paragraphs):
        if is_heading(para, 1):
            t = para.text.strip().replace(" ", "").replace("\u3000", "")
            if t == toc_display:
                continue
            first_h1_idx = i
            break

    if first_h1_idx is None or first_h1_idx == 0:
        return False

    for para in doc.paragraphs[:first_h1_idx]:
        t = para.text.strip()
        t_nospace = t.replace(" ", "").replace("\u3000", "")
        if t_nospace == "摘要":
            return True
        if re.match(cn_kw_re, t):
            return True
        if re.match(en_abs_re, t):
            return True
    return False


def _find_special_display(cfg, match_text, raw=False):
    """Find special title display text for a match keyword.
    If raw=True, return the original match text when not found."""
    for st in cfg.get("special_titles", []):
        if st["match"] == match_text:
            return st["match"] if raw else st["display"]
    return match_text


def _get_special_title_map(cfg):
    """Return {normalized_match: {display, align}} dict."""
    result = {}
    for st in cfg.get("special_titles", []):
        key = st["match"].replace(" ", "").replace("\u3000", "")
        result[key] = st
    return result


# ---------------------------------------------------------------------------
# TOC
# ---------------------------------------------------------------------------

def insert_toc(doc, cfg):
    toc_match = _find_special_display(cfg, "目录", raw=True)
    toc_depth = cfg["toc"]["depth"]
    h1_font = cfg["fonts"]["h1"]
    h1_sz_hp = str(int(cfg["sizes"]["h1"] * 2))
    toc_cfg = cfg["toc"]
    toc_font = toc_cfg.get("font", cfg["fonts"]["body"])
    toc_sz_hp = str(int(toc_cfg.get("font_size", cfg["sizes"]["body"]) * 2))
    toc_ls = toc_cfg.get("line_spacing", cfg["body"]["line_spacing"])
    toc_ls_twips = str(int(toc_ls * 240))
    latin = cfg["fonts"]["latin"]

    first_h1_el = None
    for para in doc.paragraphs:
        if para.style and para.style.name == "Heading 1":
            t = para.text.strip().replace(" ", "").replace("\u3000", "")
            if t == toc_match:
                para._element.getparent().remove(para._element)
                continue
            first_h1_el = para._element
            break

    if first_h1_el is None:
        return

    body = doc.element.body

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for sdt in list(body.findall("w:sdt", ns)):
        body.remove(sdt)

    # TOC title
    toc_display = _find_special_display(cfg, "目录")
    toc_title = OxmlElement("w:p")
    toc_title_ppr = OxmlElement("w:pPr")
    toc_title_jc = OxmlElement("w:jc")
    toc_title_jc.set(qn("w:val"), "center")
    toc_title_ppr.append(toc_title_jc)
    toc_title_spacing = OxmlElement("w:spacing")
    toc_title_spacing.set(qn("w:before"), "0")
    toc_title_spacing.set(qn("w:after"), "0")
    toc_title_spacing.set(qn("w:line"), "360")
    toc_title_spacing.set(qn("w:lineRule"), "auto")
    toc_title_ppr.append(toc_title_spacing)
    toc_title.append(toc_title_ppr)

    toc_title_run = OxmlElement("w:r")
    toc_title_rpr = OxmlElement("w:rPr")
    toc_title_rfonts = OxmlElement("w:rFonts")
    toc_title_rfonts.set(qn("w:ascii"), latin)
    toc_title_rfonts.set(qn("w:hAnsi"), latin)
    toc_title_rfonts.set(qn("w:eastAsia"), h1_font)
    toc_title_rpr.append(toc_title_rfonts)
    toc_title_sz = OxmlElement("w:sz")
    toc_title_sz.set(qn("w:val"), h1_sz_hp)
    toc_title_rpr.append(toc_title_sz)
    toc_title_szCs = OxmlElement("w:szCs")
    toc_title_szCs.set(qn("w:val"), h1_sz_hp)
    toc_title_rpr.append(toc_title_szCs)
    toc_title_bold = OxmlElement("w:b")
    toc_title_rpr.append(toc_title_bold)
    toc_title_run.append(toc_title_rpr)
    toc_title_text = OxmlElement("w:t")
    toc_title_text.set(qn("xml:space"), "preserve")
    toc_title_text.text = toc_display
    toc_title_run.append(toc_title_text)
    toc_title.append(toc_title_run)

    # TOC field
    toc_field = OxmlElement("w:p")
    toc_field_ppr = OxmlElement("w:pPr")
    toc_field_spacing = OxmlElement("w:spacing")
    toc_field_spacing.set(qn("w:line"), toc_ls_twips)
    toc_field_spacing.set(qn("w:lineRule"), "auto")
    toc_field_ppr.append(toc_field_spacing)
    toc_field.append(toc_field_ppr)

    run_begin = OxmlElement("w:r")
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run_begin.append(fld_begin)
    toc_field.append(run_begin)

    run_instr = OxmlElement("w:r")
    instr_rpr = OxmlElement("w:rPr")
    instr_rfonts = OxmlElement("w:rFonts")
    instr_rfonts.set(qn("w:ascii"), latin)
    instr_rfonts.set(qn("w:hAnsi"), latin)
    instr_rfonts.set(qn("w:eastAsia"), toc_font)
    instr_rpr.append(instr_rfonts)
    instr_sz = OxmlElement("w:sz")
    instr_sz.set(qn("w:val"), toc_sz_hp)
    instr_rpr.append(instr_sz)
    run_instr.append(instr_rpr)
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f' TOC \\o "1-{toc_depth}" \\h \\z \\u '
    run_instr.append(instr_text)
    toc_field.append(run_instr)

    run_sep = OxmlElement("w:r")
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    run_sep.append(fld_sep)
    toc_field.append(run_sep)

    run_placeholder = OxmlElement("w:r")
    ph_rpr = OxmlElement("w:rPr")
    ph_rfonts = OxmlElement("w:rFonts")
    ph_rfonts.set(qn("w:ascii"), latin)
    ph_rfonts.set(qn("w:hAnsi"), latin)
    ph_rfonts.set(qn("w:eastAsia"), toc_font)
    ph_rpr.append(ph_rfonts)
    ph_sz = OxmlElement("w:sz")
    ph_sz.set(qn("w:val"), toc_sz_hp)
    ph_rpr.append(ph_sz)
    run_placeholder.append(ph_rpr)
    ph_text = OxmlElement("w:t")
    ph_text.text = "请在 Word 中右键此处 → 更新域 → 更新整个目录"
    run_placeholder.append(ph_text)
    toc_field.append(run_placeholder)

    run_end = OxmlElement("w:r")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run_end.append(fld_end)
    toc_field.append(run_end)

    page_break = OxmlElement("w:p")
    pb_run = OxmlElement("w:r")
    br_el = OxmlElement("w:br")
    br_el.set(qn("w:type"), "page")
    pb_run.append(br_el)
    page_break.append(pb_run)

    body.insert(list(body).index(first_h1_el), page_break)
    body.insert(list(body).index(page_break), toc_field)
    body.insert(list(body).index(toc_field), toc_title)


def ensure_toc_styles(doc, cfg):
    toc_cfg = cfg["toc"]
    toc_font = toc_cfg.get("font", cfg["fonts"]["body"])
    toc_sz_hp = str(int(toc_cfg.get("font_size", cfg["sizes"]["body"]) * 2))
    toc_h1_font = toc_cfg.get("h1_font", cfg["fonts"]["h1"])
    toc_h1_sz_hp = str(int(toc_cfg.get("h1_font_size", cfg["sizes"]["h1"]) * 2))
    latin = cfg["fonts"]["latin"]

    styles_el = doc.styles.element
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    toc_depth = cfg["toc"]["depth"]

    for i in range(1, toc_depth + 1):
        style_id = f"TOC{i}"
        ea = toc_h1_font if i == 1 else toc_font
        sz_hp = toc_h1_sz_hp if i == 1 else toc_sz_hp
        found = styles_el.find(f'.//w:style[@w:styleId="{style_id}"]', ns)
        if found is not None:
            rpr = found.find("w:rPr", ns)
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                found.append(rpr)
            rfonts = rpr.find("w:rFonts", ns)
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            for theme in ["w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"]:
                if rfonts.get(qn(theme)) is not None:
                    del rfonts.attrib[qn(theme)]
            rfonts.set(qn("w:ascii"), latin)
            rfonts.set(qn("w:hAnsi"), latin)
            rfonts.set(qn("w:eastAsia"), ea)
            sz = rpr.find("w:sz", ns)
            if sz is None:
                sz = OxmlElement("w:sz")
                rpr.append(sz)
            sz.set(qn("w:val"), sz_hp)
            szCs = rpr.find("w:szCs", ns)
            if szCs is None:
                szCs = OxmlElement("w:szCs")
                rpr.append(szCs)
            szCs.set(qn("w:val"), sz_hp)
            continue

        style_el = OxmlElement("w:style")
        style_el.set(qn("w:type"), "paragraph")
        style_el.set(qn("w:styleId"), style_id)

        name_el = OxmlElement("w:name")
        name_el.set(qn("w:val"), f"toc {i}")
        style_el.append(name_el)

        based = OxmlElement("w:basedOn")
        based.set(qn("w:val"), "Normal")
        style_el.append(based)

        nxt = OxmlElement("w:next")
        nxt.set(qn("w:val"), "Normal")
        style_el.append(nxt)

        ui = OxmlElement("w:uiPriority")
        ui.set(qn("w:val"), "39")
        style_el.append(ui)

        unhide = OxmlElement("w:unhideWhenUsed")
        style_el.append(unhide)

        rpr = OxmlElement("w:rPr")
        rfonts = OxmlElement("w:rFonts")
        rfonts.set(qn("w:ascii"), latin)
        rfonts.set(qn("w:hAnsi"), latin)
        rfonts.set(qn("w:eastAsia"), ea)
        rpr.append(rfonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), sz_hp)
        rpr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), sz_hp)
        rpr.append(szCs)
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "000000")
        rpr.append(color)
        style_el.append(rpr)

        ppr = OxmlElement("w:pPr")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:line"), "360")
        spacing.set(qn("w:lineRule"), "auto")
        spacing.set(qn("w:before"), "0")
        spacing.set(qn("w:after"), "0")
        ppr.append(spacing)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:firstLine"), "0")
        if i > 1:
            ind.set(qn("w:left"), str((i - 1) * 240))
        ppr.append(ind)
        style_el.append(ppr)

        styles_el.append(style_el)


# ---------------------------------------------------------------------------
# Cover & Declaration
# ---------------------------------------------------------------------------

def _has_cover(doc, cfg, scan_limit=30):
    cover_title = cfg["cover"].get("title_text", "毕业论文")
    keywords = ["毕业论文", "毕业设计"]
    if cover_title:
        keywords.append(cover_title.replace(" ", ""))
    for para in doc.paragraphs[:scan_limit]:
        t = para.text.replace(" ", "").replace("\u3000", "")
        if any(kw in t for kw in keywords):
            return True
    return False


def insert_custom_cover(doc, cover_path):
    """Prepend a user-provided cover docx to the beginning of doc."""
    import copy as _copy
    from docx.opc.part import Part as OpcPart
    from docx.opc.packuri import PackURI

    cover_doc = Document(cover_path)

    # --- 1. Migrate styles ---
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    main_styles = doc.styles.element
    existing_ids = {
        s.get(qn("w:styleId"))
        for s in main_styles.findall(f"{{{ns_w}}}style")
    }
    for style_el in cover_doc.styles.element.findall(f"{{{ns_w}}}style"):
        sid = style_el.get(qn("w:styleId"))
        if sid and sid not in existing_ids:
            main_styles.append(_copy.deepcopy(style_el))
            existing_ids.add(sid)

    # --- 2. Migrate numbering definitions ---
    try:
        cover_num_part = cover_doc.part.numbering_part
        main_num_part = doc.part.numbering_part
        if cover_num_part is not None and main_num_part is not None:
            ns = {"w": ns_w}
            for anum in cover_num_part.element.findall("w:abstractNum", ns):
                main_num_part.element.append(_copy.deepcopy(anum))
            for num in cover_num_part.element.findall("w:num", ns):
                main_num_part.element.append(_copy.deepcopy(num))
    except Exception:
        pass

    # --- 3. Build rId mapping (images + hyperlinks) ---
    ns_a = "http://schemas.openxmlformats.org/drawingml/2006/main"
    ns_r = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    # Find next available image number
    max_img = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            m = re.search(r"image(\d+)", str(rel.target_part.partname))
            if m:
                max_img = max(max_img, int(m.group(1)))
    next_img = max_img + 1

    rid_map = {}
    for rel in cover_doc.part.rels.values():
        try:
            if "image" in rel.reltype:
                src = rel.target_part
                ext = os.path.splitext(str(src.partname))[1] or ".png"
                new_name = PackURI(f"/word/media/image{next_img}{ext}")
                next_img += 1
                img_part = OpcPart(
                    new_name, src.content_type, bytes(src.blob), doc.part.package
                )
                new_rId = doc.part.relate_to(img_part, rel.reltype)
                rid_map[rel.rId] = new_rId
            elif rel.is_external:
                # Hyperlinks and other external relationships
                new_rId = doc.part.rels.get_or_add_ext_rel(
                    rel.reltype, rel.target_ref
                )
                rid_map[rel.rId] = new_rId
        except Exception:
            pass

    # --- 4. Tags to strip (unmigrated part references) ---
    _strip_tags = {
        qn("w:commentRangeStart"), qn("w:commentRangeEnd"),
        qn("w:commentReference"), qn("w:annotationRef"),
        qn("w:footnoteReference"), qn("w:endnoteReference"),
    }

    # --- 5. Copy body elements in forward order ---
    main_body = doc.element.body
    first_child = main_body[0] if len(main_body) > 0 else None
    # Cache insertion index once; increment as we insert
    insert_idx = list(main_body).index(first_child) if first_child is not None else len(main_body)

    for el in list(cover_doc.element.body):
        if el.tag == qn("w:sectPr"):
            continue
        el_copy = _copy.deepcopy(el)

        # Strip unmigrated reference markers
        for bad in list(el_copy.iter()):
            if bad.tag in _strip_tags:
                parent = bad.getparent()
                if parent is not None:
                    parent.remove(bad)

        # Fix ALL rId references in the copied element
        for node in el_copy.iter():
            for attr_name, attr_val in list(node.attrib.items()):
                if attr_val in rid_map:
                    node.set(attr_name, rid_map[attr_val])

        main_body.insert(insert_idx, el_copy)
        insert_idx += 1

    # Page break after cover
    brk_p = OxmlElement("w:p")
    brk_r = OxmlElement("w:r")
    brk_br = OxmlElement("w:br")
    brk_br.set(qn("w:type"), "page")
    brk_r.append(brk_br)
    brk_p.append(brk_r)
    main_body.insert(insert_idx, brk_p)


def insert_cover_and_declaration(doc, cfg, config_path=None, skip_cover=False):
    cover = cfg["cover"]
    latin = cfg["fonts"]["latin"]

    def mk_run(text, ea="宋体", sz_hp=None, bold=False, uline=False):
        r = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts")
        rf.set(qn("w:ascii"), latin)
        rf.set(qn("w:hAnsi"), latin)
        rf.set(qn("w:eastAsia"), ea)
        rPr.append(rf)
        if sz_hp:
            for tag in ("w:sz", "w:szCs"):
                s = OxmlElement(tag)
                s.set(qn("w:val"), str(sz_hp))
                rPr.append(s)
        if bold:
            rPr.append(OxmlElement("w:b"))
        if uline:
            u = OxmlElement("w:u")
            u.set(qn("w:val"), "single")
            rPr.append(u)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        return r

    def mk_para(runs=None, align=None, fi=None, fi_chars=None,
                ls_auto=None, ls_exact=None, sb=None, sa=None):
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        if align:
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), align)
            pPr.append(jc)
        sp = OxmlElement("w:spacing")
        need_sp = False
        if ls_exact is not None:
            sp.set(qn("w:line"), str(ls_exact))
            sp.set(qn("w:lineRule"), "exact")
            need_sp = True
        elif ls_auto is not None:
            sp.set(qn("w:line"), str(ls_auto))
            sp.set(qn("w:lineRule"), "auto")
            need_sp = True
        if sb is not None:
            sp.set(qn("w:before"), str(sb))
            need_sp = True
        if sa is not None:
            sp.set(qn("w:after"), str(sa))
            need_sp = True
        if need_sp:
            pPr.append(sp)
        if fi is not None or fi_chars is not None:
            ind = OxmlElement("w:ind")
            if fi is not None:
                ind.set(qn("w:firstLine"), str(fi))
            if fi_chars is not None:
                ind.set(qn("w:firstLineChars"), str(fi_chars))
            pPr.append(ind)
        p.append(pPr)
        for r in (runs or []):
            p.append(r)
        return p

    NBSP = "\u00a0"  # 不换行空格 (下划线填充用，外观与普通空格相同)

    def mk_field(label, uline_chars=33):
        return mk_para([
            mk_run(label, sz_hp=30, bold=True),
            mk_run(" ", sz_hp=30),
            mk_run(NBSP * uline_chars, sz_hp=30, uline=True),
        ], fi_chars=400, ls_exact=700, sa=0, sb=0)

    elements = []

    if not skip_cover:
        # Cover page
        title_sz_hp = int(cover["title_font_size"] * 2)
        thesis_sz_hp = int(cover["thesis_title_size"] * 2)
        thesis_font = cover["thesis_title_font"]

        elements.append(mk_para(align="center", ls_auto=360))
        elements.append(mk_para(
            [mk_run(cover["title_text"], sz_hp=title_sz_hp, bold=True)],
            align="center", sa=161,
        ))
        elements.append(mk_para(align="center", ls_auto=360))
        elements.append(mk_para(
            [mk_run(cover["thesis_title_placeholder"], ea=thesis_font,
                    sz_hp=thesis_sz_hp, bold=True)],
            align="center", sa=161,
        ))
        for sz in (30, 30, 21, 21, 30, 30):
            elements.append(mk_para([mk_run(" ", sz_hp=sz)], ls_auto=360, sa=0, sb=0))

        for field in cover["fields"]:
            elements.append(mk_field(field["label"], field["underline_chars"]))

        adv = cover["advisor"]
        elements.append(mk_para([
            mk_run(adv["label"], sz_hp=30, bold=True),
            mk_run(" ", sz_hp=30),
            mk_run(NBSP * adv["underline_chars"], sz_hp=30, uline=True),
            mk_run(" ", sz_hp=30),
            mk_run(adv["title_label"], sz_hp=30, bold=True),
            mk_run(" ", sz_hp=30),
            mk_run(NBSP * adv["title_underline_chars"], sz_hp=30, uline=True),
        ], fi_chars=400, ls_exact=700, sa=0, sb=0))

        dt = cover["date"]
        date_runs = [mk_run(dt["label"], sz_hp=30, bold=True)]
        for seg in dt["segments"]:
            date_runs.extend([
                mk_run(" ", sz_hp=30),
                mk_run(NBSP * dt["segment_underline_chars"], sz_hp=30, uline=True),
                mk_run(" ", sz_hp=30),
                mk_run(seg, sz_hp=30, bold=True),
            ])
        elements.append(mk_para(date_runs, fi_chars=400, ls_exact=700, sa=0, sb=0))

        elements.append(mk_para(align="center", sa=161))

        # Page break after cover
        pb = OxmlElement("w:p")
        pb_r = OxmlElement("w:r")
        pb_br = OxmlElement("w:br")
        pb_br.set(qn("w:type"), "page")
        pb_r.append(pb_br)
        pb.append(pb_r)
        elements.append(pb)

    # Declaration pages
    declarations = cfg.get("declarations", [])
    if declarations:
        elements.append(mk_para(ls_auto=360))

    for idx, decl in enumerate(declarations):
        h1_font = cfg["fonts"]["h1"]
        elements.append(mk_para(
            [mk_run(decl["title"], ea=h1_font, sz_hp=28)],
            align="center", sb=161, sa=161, ls_auto=360,
        ))
        elements.append(mk_para([mk_run(
            decl["body"], sz_hp=24,
        )], fi=540, ls_auto=360))
        elements.append(mk_para(fi=540, ls_auto=360))
        elements.append(mk_para([mk_run(
            decl["signature"], sz_hp=24,
        )], fi=540, ls_auto=360))

        if "date_line" in decl:
            last_el = mk_para([mk_run(
                decl["date_line"], sz_hp=24,
            )], fi=540, ls_auto=360)
            elements.append(last_el)
        else:
            last_el = elements[-1]

        if idx < len(declarations) - 1:
            for _ in range(4):
                elements.append(mk_para(ls_auto=360))
        elif idx == len(declarations) - 1:
            elements.append(mk_para(fi=540, ls_auto=360))

    # Section break after declarations
    if elements:
        final_el = elements[-1]
        pPr = final_el.find(qn("w:pPr"))
        if pPr is None:
            pPr = OxmlElement("w:pPr")
            final_el.insert(0, pPr)
        sect_pr = OxmlElement("w:sectPr")
        existing_sect = doc.sections[0]._sectPr
        for attr_name in ("pgSz", "pgMar"):
            src = existing_sect.find(qn(f"w:{attr_name}"))
            if src is not None:
                sect_pr.append(copy.deepcopy(src))
        pPr.append(sect_pr)

    body = doc.element.body
    for el in reversed(elements):
        body.insert(0, el)

    # Clear footer for cover/declarations section (no page numbers)
    if len(doc.sections) > 0:
        s0_footer = doc.sections[0].footer
        s0_footer.is_linked_to_previous = False
        for p in s0_footer.paragraphs:
            p.clear()
    if len(doc.sections) > 1:
        doc.sections[1].footer.is_linked_to_previous = False

    # Logo
    if not skip_cover:
        logo_path = resolve_logo_path(cfg, config_path)
        if logo_path:
            p0_para = doc.paragraphs[0]
            run = p0_para.add_run()
            run.add_picture(logo_path,
                            width=Pt(cover["logo_width_pt"]),
                            height=Pt(cover["logo_height_pt"]))

    if len(doc.sections) > 1:
        doc.sections[1].footer.is_linked_to_previous = False


# ---------------------------------------------------------------------------
# Structure validation (config-driven)
# ---------------------------------------------------------------------------

def validate_structure(doc, cfg):
    warnings = []
    paras = doc.paragraphs
    texts = [p.text.strip() for p in paras]
    texts_nospace = [t.replace(" ", "").replace("\u3000", "") for t in texts]

    sec = cfg["sections"]
    st_map = _get_special_title_map(cfg)

    # Front matter checks
    has_cn_abstract = any(t == "摘要" for t in texts_nospace)
    cn_kw_pat = sec.get("cn_keywords_pattern", r"关键词[：:]")
    has_cn_keywords = any(re.match(cn_kw_pat, t) for t in texts_nospace)
    en_abs_pat = sec.get("en_abstract_pattern", r"(?i)abstract[：:]?")
    has_en_abstract = any(re.match(en_abs_pat, t) for t in texts_nospace)
    en_kw_pat = sec.get("en_keywords_pattern", r"(?i)keywords?[：:]")
    has_en_keywords = any(re.match(en_kw_pat, t.replace(" ", "")) for t in texts)

    if not has_cn_abstract:
        warnings.append("缺少中文摘要标题")
    if not has_cn_keywords:
        warnings.append("缺少中文关键词")
    if not has_en_abstract:
        warnings.append("缺少英文摘要 (Abstract)")
    if not has_en_keywords:
        warnings.append("缺少英文关键词 (Key words)")

    cn_kw_idx = next((i for i, t in enumerate(texts_nospace)
                      if re.match(cn_kw_pat, t)), None)
    en_abs_idx = next((i for i, t in enumerate(texts_nospace)
                       if re.match(en_abs_pat, t)), None)
    if cn_kw_idx is not None and en_abs_idx is not None and cn_kw_idx < en_abs_idx:
        between = [texts[j] for j in range(cn_kw_idx + 1, en_abs_idx) if texts[j]]
        has_en_title = any(re.search(r"[A-Za-z]{4,}", t) for t in between)
        has_affiliation = any(re.search(r"(?i)(university|college|china|institute)", t)
                              for t in between)
        if not has_en_title:
            warnings.append("英文摘要页缺少英文题目")
        if not has_affiliation:
            warnings.append("英文摘要页缺少作者英文名与单位信息")

    # Body checks
    chapter_pat = sec.get("chapter_pattern", r"第\s*\d+\s*章")
    has_chapter_h1 = False
    ref_key = "参考文献"
    thanks_key = "致谢"
    if "参考文献" in st_map:
        ref_key = st_map["参考文献"]["match"]
    if "致谢" in st_map:
        thanks_key = st_map["致谢"]["match"]

    has_refs = any(t == ref_key.replace(" ", "").replace("\u3000", "") for t in texts_nospace)
    has_thanks = any(t == thanks_key.replace(" ", "").replace("\u3000", "") for t in texts_nospace)

    # Check for TOC
    toc_key = _find_special_display(cfg, "目录", raw=True)
    has_toc = any(t == toc_key for t in texts_nospace)
    if not has_toc:
        warnings.append("缺少「目录」标题")

    # Check figure/table captions
    cap_cfg = cfg.get("captions", {})
    fig_pat = cap_cfg.get("figure_pattern", r"^图\s*\d")
    tbl_pat = cap_cfg.get("table_pattern", r"^(续)?表\s*\d")
    has_images = any(
        el.tag.endswith("}blip") for el in doc.element.body.iter()
    )
    has_tables = len(doc.tables) > 0
    has_fig_cap = any(re.match(fig_pat, t) for t in texts)
    has_tbl_cap = any(re.match(tbl_pat, t) for t in texts)
    if has_images and not has_fig_cap:
        warnings.append("检测到插图但缺少图题（如「图1 xxx」）")
    if has_tables and not has_tbl_cap:
        warnings.append("检测到表格但缺少表题（如「表1 xxx」）")

    # Check heading styles are applied
    has_heading_styles = any(
        p.style and p.style.name in ("Heading 1", "Heading 2", "Heading 3", "Heading 4", "样式1")
        for p in paras if p.text.strip())
    if not has_heading_styles:
        warnings.append("未检测到标题样式（请确保 Word 中已对标题应用 Heading 1/2/3 样式）")

    for p in paras:
        sn = p.style.name if p.style else ""
        t = p.text.strip()
        if sn in ("Heading 1", "样式1") and re.match(chapter_pat, t):
            has_chapter_h1 = True
            break

    if not has_chapter_h1:
        warnings.append("未检测到正文章节标题")
    if not has_refs:
        warnings.append("缺少「参考文献」标题")
    if not has_thanks:
        warnings.append("缺少「致谢」标题")

    # Heading numbering checks
    appendix_pat = sec.get("appendix_pattern", r"^附录\s*[A-Z]")
    h1_pat = re.compile(f"({chapter_pat}|{appendix_pat})")
    h2_pat = re.compile(sec.get("h2_pattern", r"^\d+\.\d+\s"))
    h3_pat = re.compile(sec.get("h3_pattern", r"^\d+\.\d+\.\d+\s"))
    h4_pat = re.compile(sec.get("h4_pattern", r"^\d+\.\d+\.\d+\.\d+\s"))

    special_h1_set = set(st_map.keys())
    special_h1_set.update(s.replace(" ", "").replace("\u3000", "")
                          for s in sec.get("special_h1", []))

    for p in paras:
        sn = p.style.name if p.style else ""
        t = p.text.strip()
        t_nospace = t.replace(" ", "").replace("\u3000", "")
        if not t:
            continue

        if sn == "Heading 1" or sn == "样式1":
            if t_nospace not in special_h1_set and not h1_pat.match(t):
                warnings.append(f"一级标题缺少编号: \"{t}\"")
        elif sn == "Heading 2":
            if not h2_pat.match(t):
                warnings.append(f"二级标题缺少编号: \"{t}\"")
        elif sn == "Heading 3":
            if not h3_pat.match(t):
                warnings.append(f"三级标题缺少编号: \"{t}\"")
        elif sn == "Heading 4":
            if not h4_pat.match(t):
                warnings.append(f"四级标题缺少编号: \"{t}\"")

    if warnings:
        print("=" * 50, file=sys.stderr)
        print("结构检查警告:", file=sys.stderr)
        for w in warnings:
            print(f"  ⚠ {w}", file=sys.stderr)
        print("=" * 50, file=sys.stderr)

    return warnings


# ---------------------------------------------------------------------------
# Heading auto-renumbering
# ---------------------------------------------------------------------------

_CN_NUMERALS = "零一二三四五六七八九十"


def _int_to_cn(n):
    """Convert integer (1-99) to Chinese numeral."""
    if n <= 10:
        return _CN_NUMERALS[n]
    if n < 20:
        return "十" + (_CN_NUMERALS[n - 10] if n > 10 else "")
    tens, ones = divmod(n, 10)
    return (_CN_NUMERALS[tens] if tens > 1 else "") + "十" + \
           (_CN_NUMERALS[ones] if ones else "")


def _renumber_h1_text(text, new_num, pattern):
    """Replace the chapter number in an H1 heading text."""
    # "第X章" style
    if re.search(r"第\s*\d+\s*章", text):
        return re.sub(r"(第\s*)\d+(\s*章)", fr"\g<1>{new_num}\2", text, count=1)
    # "Chapter X" style
    if re.search(r"(?i)Chapter\s+\d+", text):
        return re.sub(r"(?i)(Chapter\s+)\d+", fr"\g<1>{new_num}", text, count=1)
    # Chinese numeral style "一、绪论"
    if re.match(r"^[一二三四五六七八九十百]+、", text):
        cn = _int_to_cn(new_num)
        return re.sub(r"^[一二三四五六七八九十百]+", cn, text, count=1)
    # Plain number style "X title"
    if re.match(r"^\d+\s", text):
        return re.sub(r"^\d+", str(new_num), text, count=1)
    return text


def _renumber_sub_text(text, prefix):
    """Replace X.Y[.Z[.W]] or Chinese numeral sub-heading prefix."""
    # Chinese numeral H2: （一）xxx
    if re.match(r"^（[一二三四五六七八九十百]+）", text):
        return re.sub(r"^（[一二三四五六七八九十百]+）", prefix, text, count=1)
    return re.sub(r"^[\d.]+", prefix, text, count=1)


def renumber_headings(doc, cfg):
    """Detect heading numbering gaps and auto-fix. Returns list of change descriptions."""
    sec = cfg.get("sections", {})
    chap_pat = re.compile(sec.get("chapter_pattern", r"^第\s*\d+\s*章\b"))
    appendix_pat = re.compile(sec.get("appendix_pattern", r"^附录\s*[A-Z]"))
    h2_pat = re.compile(sec.get("h2_pattern", r"^\d+\.\d+\s"))
    h3_pat = re.compile(sec.get("h3_pattern", r"^\d+\.\d+\.\d+\s"))
    h4_pat = re.compile(sec.get("h4_pattern", r"^\d+\.\d+\.\d+\.\d+\s"))

    st_map = _get_special_title_map(cfg)
    special_set = set(st_map.keys())
    special_set.update(s.replace(" ", "").replace("\u3000", "")
                       for s in sec.get("special_h1", []))

    changes = []
    chap_n = 0
    h2_n = h3_n = h4_n = 0
    in_appendix = False

    for para in doc.paragraphs:
        sn = para.style.name if para.style else ""
        t = para.text.strip()
        t_nospace = t.replace(" ", "").replace("\u3000", "")
        if not t:
            continue

        if sn in ("Heading 1", "样式1"):
            if t_nospace in special_set:
                continue
            if appendix_pat.match(t):
                in_appendix = True
                continue
            if in_appendix:
                continue
            if chap_pat.match(t):
                chap_n += 1
                h2_n = h3_n = h4_n = 0
                new_t = _renumber_h1_text(t, chap_n, chap_pat.pattern)
                if new_t != t:
                    changes.append(f"  H1: \"{t}\" → \"{new_t}\"")
                    _replace_para_text(para, new_t)
        elif sn == "Heading 2" and not in_appendix:
            if h2_pat.match(t):
                h2_n += 1
                h3_n = h4_n = 0
                new_t = _renumber_sub_text(t, f"{chap_n}.{h2_n}")
                if new_t != t:
                    changes.append(f"  H2: \"{t}\" → \"{new_t}\"")
                    _replace_para_text(para, new_t)
        elif sn == "Heading 3" and not in_appendix:
            if h3_pat.match(t):
                h3_n += 1
                h4_n = 0
                new_t = _renumber_sub_text(t, f"{chap_n}.{h2_n}.{h3_n}")
                if new_t != t:
                    changes.append(f"  H3: \"{t}\" → \"{new_t}\"")
                    _replace_para_text(para, new_t)
        elif sn == "Heading 4" and not in_appendix:
            if h4_pat.match(t):
                h4_n += 1
                new_t = _renumber_sub_text(t, f"{chap_n}.{h2_n}.{h3_n}.{h4_n}")
                if new_t != t:
                    changes.append(f"  H4: \"{t}\" → \"{new_t}\"")
                    _replace_para_text(para, new_t)

    return changes


def _replace_para_text(para, new_text):
    """Replace paragraph text while preserving the first run's formatting."""
    if para.runs:
        first_run = para.runs[0]
        font_props = {
            "name": first_run.font.name,
            "size": first_run.font.size,
            "bold": first_run.font.bold,
        }
        ea = first_run.font.element.find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
        ea_name = ea.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia") if ea is not None else None
    else:
        font_props = None
        ea_name = None

    para.text = new_text

    if font_props and para.runs:
        r = para.runs[0]
        r.font.name = font_props["name"]
        r.font.size = font_props["size"]
        r.font.bold = font_props["bold"]
        if ea_name:
            from docx.oxml.ns import qn as _qn
            rFonts = r.font.element.find(_qn("w:rFonts"))
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement("w:rFonts")
                r.font.element.insert(0, rFonts)
            rFonts.set(_qn("w:eastAsia"), ea_name)


def normalize_heading_spacing(doc, cfg):
    """Normalize spacing between heading number and title to 1字距 (全角空格)."""
    sec = cfg.get("sections", {})
    chap_re = re.compile(sec.get("chapter_pattern", r"^第\s*\d+\s*章\b"))
    st_map = _get_special_title_map(cfg)
    JIJU = "  "  # 1字距 = 2个普通空格

    for para in doc.paragraphs:
        sn = para.style.name if para.style else ""
        t = para.text.strip()
        if not t:
            continue
        t_nospace = t.replace(" ", "").replace("\u3000", "")

        new_t = None
        if sn in ("Heading 1", "样式1"):
            if t_nospace in st_map:
                continue
            # "第X章 xxx"
            m = re.match(r"(第\s*\d+\s*章)\s*(.*)", t)
            if m and m.group(2):
                new_t = m.group(1) + JIJU + m.group(2)
            # "Chapter X xxx"
            if new_t is None:
                m = re.match(r"((?i)Chapter\s+\d+)\s+(.*)", t)
                if m and m.group(2):
                    new_t = m.group(1) + JIJU + m.group(2)
            # "一、xxx"
            if new_t is None:
                m = re.match(r"([一二三四五六七八九十百]+、)\s*(.*)", t)
                if m and m.group(2):
                    new_t = m.group(1) + JIJU + m.group(2)
            # "1 xxx" (pure number)
            if new_t is None:
                m = re.match(r"(\d+)\s+(.*)", t)
                if m and m.group(2):
                    new_t = m.group(1) + JIJU + m.group(2)
        elif sn in ("Heading 2", "Heading 3", "Heading 4"):
            # "（一）xxx"
            m = re.match(r"(（[一二三四五六七八九十百]+）)\s*(.*)", t)
            if m and m.group(2):
                new_t = m.group(1) + JIJU + m.group(2)
            # numeric: "1.1 xxx", "1.1.1 xxx", "1. xxx"
            if new_t is None:
                m = re.match(r"([\d.]+)\s*(.*)", t)
                if m and m.group(2):
                    new_t = m.group(1) + JIJU + m.group(2)
            # "(1)xxx"
            if new_t is None:
                m = re.match(r"(\(\d+\))\s*(.*)", t)
                if m and m.group(2):
                    new_t = m.group(1) + JIJU + m.group(2)

        if new_t is not None and new_t != t:
            _replace_para_text(para, new_t)


# ---------------------------------------------------------------------------
# Auto-assign heading styles for docx input
# ---------------------------------------------------------------------------

_SENTENCE_ENDINGS = set("。！？；")


def auto_assign_heading_styles(doc, cfg):
    """Detect headings in Normal-style paragraphs and assign Heading styles."""
    sec = cfg.get("sections", {})
    chap_re = re.compile(sec.get("chapter_pattern", r"^第\s*\d+\s*章\b"))
    appendix_re = re.compile(sec.get("appendix_pattern", r"^附录\s*[A-Z]"))
    h2_re = re.compile(sec.get("h2_pattern", r"^\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    h3_re = re.compile(sec.get("h3_pattern", r"^\d+\.\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    h4_re = re.compile(sec.get("h4_pattern", r"^\d+\.\d+\.\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))

    st_map = _get_special_title_map(cfg)
    special_h1_set = set()
    for s in sec.get("special_h1", []):
        special_h1_set.add(s.replace(" ", "").replace("\u3000", ""))
    special_h1_set.update(st_map.keys())

    _skip_styles = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "样式1"}

    changes = []
    for para in doc.paragraphs:
        sn = para.style.name if para.style else ""
        if sn in _skip_styles:
            continue
        t = para.text.strip()
        if not t:
            continue
        t_nospace = t.replace(" ", "").replace("\u3000", "")

        target = None

        # H1: 第X章
        if chap_re.match(t):
            target = "Heading 1"
        # H1: special titles (摘要, 参考文献, 致谢, etc.)
        elif t_nospace in special_h1_set:
            target = "Heading 1"
        # H1: 附录X
        elif appendix_re.match(t):
            target = "Heading 1"
        # H1: pure number "1 绪论" or "2文献综述"
        elif re.match(r"^\d+(\s|(?=[\u4e00-\u9fff]))", t) and not re.match(r"^\d+\.\d+", t):
            if len(t) <= 50 and t[-1] not in _SENTENCE_ENDINGS:
                target = "Heading 1"
        else:
            # H4/H3/H2: short title guard
            if len(t) <= 50 and t[-1] not in _SENTENCE_ENDINGS:
                if h4_re.match(t):
                    target = "Heading 4"
                elif h3_re.match(t):
                    target = "Heading 3"
                elif h2_re.match(t):
                    target = "Heading 2"

        if target and target in [s.name for s in doc.styles]:
            para.style = doc.styles[target]
            changes.append(f"  {target}: \"{t}\"")

    return changes

_CITE_NUM_RE = re.compile(r'\[(\d+(?:\s*[,，\-–]\s*\d+)*)\]')
_CITE_AY_OUTER = re.compile(r'[（(](.+?)[）)]')
_CITE_AY_INNER = re.compile(r'(.+?)[,，]\s*((?:19|20)\d{2}[a-z]?)\s*$')
_REF_NUM_RE = re.compile(r'^\[(\d+)\]\s*')
_REF_TYPE_RE = re.compile(r'\[([A-Z]{1,2}(?:/[A-Z]{1,2})?)\]')
_REF_YEAR_RE = re.compile(r'(?:19|20)\d{2}[a-z]?')
_GBT_VALID_TYPES = {
    "J", "M", "C", "D", "R", "S", "P", "A", "Z", "N",
    "EB/OL", "OL", "DB/OL", "CP/DK", "DB", "CP",
}


def _parse_cite_numbers(inner):
    """Parse inner text of [N,M-K] into a list of integers."""
    nums = []
    for part in re.split(r'[,，]', inner):
        part = part.strip()
        rm = re.match(r'(\d+)\s*[-–]\s*(\d+)', part)
        if rm:
            nums.extend(range(int(rm.group(1)), int(rm.group(2)) + 1))
        elif re.match(r'\d+$', part):
            nums.append(int(part))
    return nums


def _extract_primary_author(author_str):
    """Extract first author from composite author string."""
    return re.split(
        r'等|[和与&,，]|\s+and\s+|\s+et\s+al', author_str, maxsplit=1
    )[0].strip()


def check_citations(doc, cfg):
    """Check citation-reference cross-matching. Returns list of warnings."""
    warnings = []
    sec = cfg.get("sections", {})
    st_map = _get_special_title_map(cfg)

    ref_key = "参考文献"
    if "参考文献" in st_map:
        ref_key = st_map["参考文献"]["match"]
    ref_key_norm = ref_key.replace(" ", "").replace("\u3000", "")

    chap_pat = re.compile(sec.get("chapter_pattern", r"^第\s*\d+\s*章"))

    # --- locate reference section & body range ---
    paras = doc.paragraphs
    ref_start = ref_end = body_start = None

    # collect special title norms for boundary detection
    _boundary_norms = set()
    for st in sec.get("special_titles", []):
        n = st["match"].replace(" ", "").replace("\u3000", "")
        if n != ref_key_norm:
            _boundary_norms.add(n)
    _ap = sec.get("appendix_pattern", r"^附录\s*[A-Z]?")
    if _ap.endswith("[A-Z]"):
        _ap += "?"
    appendix_re = re.compile(_ap)

    for i, p in enumerate(paras):
        sn = p.style.name if p.style else ""
        t_strip = p.text.strip()
        t_norm = t_strip.replace(" ", "").replace("\u3000", "")
        is_h1 = sn in ("Heading 1", "样式1")

        if is_h1 and body_start is None and chap_pat.match(t_strip):
            body_start = i
        if is_h1 and t_norm == ref_key_norm:
            ref_start = i + 1
        elif ref_start is not None and ref_end is None:
            # end reference section at any heading OR known boundary title
            if is_h1 or (sn.startswith("Heading") and (
                    t_norm in _boundary_norms or appendix_re.match(t_strip))):
                ref_end = i

    if ref_start is None:
        return []
    if ref_end is None:
        ref_end = len(paras)
    if body_start is None:
        body_start = 0

    # --- parse reference entries ---
    ref_entries = []
    for i in range(ref_start, ref_end):
        p = paras[i]
        sn = p.style.name if p.style else ""
        t = p.text.strip()
        t_norm = t.replace(" ", "").replace("\u3000", "")

        # stop at section boundaries (heading / special title / appendix)
        if sn and sn.startswith("Heading"):
            break
        if t_norm in _boundary_norms or appendix_re.match(t):
            break
        if not t:
            continue

        entry = {"text": t, "idx": i}

        m = _REF_NUM_RE.match(t)
        entry["num"] = int(m.group(1)) if m else None
        t_body = t[m.end():] if m else t

        tm = _REF_TYPE_RE.search(t)
        entry["type"] = tm.group(1) if tm else None

        years = _REF_YEAR_RE.findall(t)
        entry["year"] = years[0] if years else None

        am = re.match(r'(.+?(?:\.[A-Z]\.)*)\.\s*(?=[^A-Z])', t_body)
        if not am:
            am = re.match(r'(.+?)．', t_body)
        entry["authors"] = am.group(1).strip() if am else t_body[:30].strip()

        ref_entries.append(entry)

    if not ref_entries:
        return []

    # --- scan body citations (skip appendix) ---
    num_cites = []   # (number, para_index)
    ay_cites = []    # (author, year, para_index)
    in_appendix = False

    for i in range(body_start, ref_start - 1):
        p = paras[i]
        sn = p.style.name if p.style else ""
        t_strip = p.text.strip()

        if sn in ("Heading 1", "样式1"):
            in_appendix = bool(appendix_re.match(t_strip))
        if sn.startswith("Heading") or in_appendix:
            continue
        if not t_strip:
            continue

        for m in _CITE_NUM_RE.finditer(t_strip):
            for n in _parse_cite_numbers(m.group(1)):
                num_cites.append((n, i))

        for m in _CITE_AY_OUTER.finditer(t_strip):
            inner = m.group(1)
            for seg in re.split(r'[;；]', inner):
                seg = seg.strip()
                am = _CITE_AY_INNER.match(seg)
                if am:
                    author = am.group(1).strip()
                    if re.fullmatch(r'[\d\s\-–—年]+', author):
                        continue  # skip year ranges like "2015—2022年"
                    ay_cites.append((author, am.group(2).strip(), i))

    # --- auto-detect dominant style ---
    style = "numbered" if len(num_cites) >= len(ay_cites) else "author-year"

    # --- cross-match ---
    if style == "numbered":
        ref_nums = {e["num"]: e for e in ref_entries if e["num"] is not None}

        # reference list ordering
        nums_list = [e["num"] for e in ref_entries if e["num"] is not None]
        if nums_list:
            expected = list(range(nums_list[0], nums_list[0] + len(nums_list)))
            if nums_list != expected:
                gaps = sorted(set(expected) - set(nums_list))
                if gaps:
                    warnings.append(f"参考文献编号不连续，缺少: {gaps}")
            seen = set()
            for n in nums_list:
                if n in seen:
                    warnings.append(f"参考文献编号重复: [{n}]")
                seen.add(n)

        # body citation first-appearance order
        first_seen = []
        for n, _ in num_cites:
            if n not in first_seen:
                first_seen.append(n)
        if first_seen and first_seen != sorted(first_seen):
            preview = first_seen[:15]
            warnings.append(
                f"正文引用编号未按首次出现顺序排列"
                f"（前{len(preview)}个: {preview}）"
            )

        # unmatched
        cited_set = {n for n, _ in num_cites}
        ref_set = set(ref_nums.keys())
        diff_cite = sorted(cited_set - ref_set)
        diff_ref = sorted(ref_set - cited_set)
        if diff_cite:
            warnings.append(f"正文引用了但文末无对应条目: {diff_cite}")
        if diff_ref:
            warnings.append(f"文末有条目但正文未引用: {diff_ref}")

    else:  # author-year
        unmatched = []
        for author_str, year_str, _ in ay_cites:
            primary = _extract_primary_author(author_str)
            found = any(
                e["year"] and e["year"][:4] == year_str[:4]
                and primary and primary in e["authors"]
                for e in ref_entries
            )
            if not found:
                tag = f"（{author_str}，{year_str}）"
                if tag not in unmatched:
                    unmatched.append(tag)
        if unmatched:
            warnings.append(
                f"正文引用了但文末无匹配条目: {', '.join(unmatched[:15])}"
            )

        ref_ay = set()
        for e in ref_entries:
            if e["year"] and e["authors"]:
                ref_ay.add((_extract_primary_author(e["authors"]), e["year"][:4]))
        cited_ay = set()
        for a, y, _ in ay_cites:
            cited_ay.add((_extract_primary_author(a), y[:4]))
        uncited = ref_ay - cited_ay
        if uncited:
            tags = [f"{a}({y})" for a, y in sorted(uncited)]
            warnings.append(f"文末有条目但正文未引用: {', '.join(tags[:15])}")

    # --- GB/T 7714 format check ---
    for e in ref_entries:
        if e["type"] is None:
            warnings.append(f"参考文献缺少类型标识[J]/[M]/..: \"{e['text'][:50]}\"")
        elif e["type"] not in _GBT_VALID_TYPES:
            warnings.append(f"参考文献类型标识不规范[{e['type']}]: \"{e['text'][:50]}\"")
        if not e["year"]:
            warnings.append(f"参考文献缺少年份: \"{e['text'][:50]}\"")

    return warnings


# ---------------------------------------------------------------------------
# Cross-reference fields for citations
# ---------------------------------------------------------------------------

def _make_text_run_el(text, rPr_el=None):
    """Create a <w:r> element with text and optional formatting."""
    r = OxmlElement('w:r')
    if rPr_el is not None:
        r.append(copy.deepcopy(rPr_el))
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def _make_field_runs(instr, display, rPr_el=None):
    """Create Word field XML: begin + instrText + separate + display + end."""
    els = []
    for ftype in ('begin', None, 'separate', None, 'end'):
        r = OxmlElement('w:r')
        if rPr_el is not None:
            r.append(copy.deepcopy(rPr_el))
        if ftype in ('begin', 'separate', 'end'):
            fc = OxmlElement('w:fldChar')
            fc.set(qn('w:fldCharType'), ftype)
            r.append(fc)
        elif len(els) == 1:  # instrText (second element)
            it = OxmlElement('w:instrText')
            it.set(qn('xml:space'), 'preserve')
            it.text = f' {instr} '
            r.append(it)
        else:  # display text (fourth element)
            t = OxmlElement('w:t')
            t.set(qn('xml:space'), 'preserve')
            t.text = display
            r.append(t)
        els.append(r)
    return els


def _parse_cite_structure(inner):
    """Parse [N,M-K] inner text preserving structure.
    Returns list of ('num', N) | ('range', (start, end)) | ('sep', ',').
    """
    parts = []
    for seg in re.split(r'([,，])', inner):
        seg = seg.strip()
        if seg in (',', '，'):
            if parts:
                parts.append(('sep', ','))
            continue
        rm = re.match(r'(\d+)\s*[-–]\s*(\d+)', seg)
        if rm:
            parts.append(('range', (int(rm.group(1)), int(rm.group(2)))))
        elif re.match(r'\d+$', seg):
            parts.append(('num', int(seg)))
    return parts


def _append_char_segment(p_el, chars):
    """Append chars to paragraph XML, grouping consecutive same-formatting."""
    if not chars:
        return
    cur_rPr = chars[0][1]
    cur_text = ""
    for ch, rPr in chars:
        if rPr is cur_rPr:
            cur_text += ch
        else:
            if cur_text:
                p_el.append(_make_text_run_el(cur_text, cur_rPr))
            cur_rPr = rPr
            cur_text = ch
    if cur_text:
        p_el.append(_make_text_run_el(cur_text, cur_rPr))


def _apply_ref_crosslinks(doc, cfg):
    """Replace static [N] with SEQ/REF Word fields for cross-referencing."""
    sec = cfg.get("sections", {})
    st_map = _get_special_title_map(cfg)

    ref_key_norm = "参考文献"
    if "参考文献" in st_map:
        ref_key_norm = st_map["参考文献"]["match"].replace(" ", "").replace("\u3000", "")

    chap_pat = re.compile(sec.get("chapter_pattern", r"^第\s*\d+\s*章"))
    _ap = sec.get("appendix_pattern", r"^附录\s*[A-Z]?")
    if _ap.endswith("[A-Z]"):
        _ap += "?"
    appendix_re = re.compile(_ap)

    _boundary_norms = set()
    for st in sec.get("special_titles", []):
        n = st["match"].replace(" ", "").replace("\u3000", "")
        if n != ref_key_norm:
            _boundary_norms.add(n)

    paras = doc.paragraphs
    ref_start = ref_end = body_start = None

    for i, p in enumerate(paras):
        sn = p.style.name if p.style else ""
        t_strip = p.text.strip()
        t_norm = t_strip.replace(" ", "").replace("\u3000", "")
        is_h1 = sn in ("Heading 1", "样式1")
        if is_h1 and body_start is None and chap_pat.match(t_strip):
            body_start = i
        if is_h1 and t_norm == ref_key_norm:
            ref_start = i + 1
        elif ref_start is not None and ref_end is None:
            if is_h1 or (sn.startswith("Heading") and (
                    t_norm in _boundary_norms or appendix_re.match(t_strip))):
                ref_end = i

    if ref_start is None:
        return
    if ref_end is None:
        ref_end = len(paras)
    if body_start is None:
        body_start = 0

    # --- detect citation style ---
    num_count = ay_count = 0
    for i in range(body_start, ref_start - 1):
        t = paras[i].text
        num_count += len(_CITE_NUM_RE.findall(t))
        for m in _CITE_AY_OUTER.finditer(t):
            inner = m.group(1)
            for seg in re.split(r'[;；]', inner):
                if _CITE_AY_INNER.match(seg.strip()):
                    ay_count += 1
    is_numbered = not (ay_count > 0 and num_count == 0 and ay_count > num_count)

    # --- Step 1: reference entries → SEQ fields + bookmarks (always) ---
    bm_id = 1000
    bookmark_map = {}  # {original_num: bookmark_name}

    for i in range(ref_start, ref_end):
        p = paras[i]
        sn = p.style.name if p.style else ""
        t = p.text.strip()
        if sn and sn.startswith("Heading"):
            break
        t_norm = t.replace(" ", "").replace("\u3000", "")
        if t_norm in _boundary_norms or appendix_re.match(t):
            break
        if not t:
            continue

        m = _REF_NUM_RE.match(t)
        if not m:
            continue

        num = int(m.group(1))
        bm_name = f"_Ref{num}"
        bookmark_map[num] = bm_name

        # build char-level formatting map
        p_el = p._element
        runs = list(p.runs)
        if not runs:
            continue
        chars = []
        for r in runs:
            r_rPr = r._element.find(qn('w:rPr'))
            for ch in (r.text or ""):
                chars.append((ch, r_rPr))
        rPr0 = chars[0][1] if chars else None
        prefix_end = m.end()

        # clear runs (keep pPr)
        for child in list(p_el):
            if child.tag != qn('w:pPr'):
                p_el.remove(child)

        # rebuild: [<bookmark SEQ>] remaining
        p_el.append(_make_text_run_el('[', rPr0))
        bm_start = OxmlElement('w:bookmarkStart')
        bm_start.set(qn('w:id'), str(bm_id))
        bm_start.set(qn('w:name'), bm_name)
        p_el.append(bm_start)
        for fel in _make_field_runs('SEQ Ref', str(num), rPr0):
            p_el.append(fel)
        bm_end = OxmlElement('w:bookmarkEnd')
        bm_end.set(qn('w:id'), str(bm_id))
        p_el.append(bm_end)
        p_el.append(_make_text_run_el('] ', rPr0))
        _append_char_segment(p_el, chars[prefix_end:])

        bm_id += 1

    if not bookmark_map:
        return

    # --- Step 2: body citations → REF fields (numbered style only) ---
    if not is_numbered or num_count == 0:
        return
    in_appendix = False
    for i in range(body_start, ref_start - 1):
        p = paras[i]
        sn = p.style.name if p.style else ""
        t_strip = p.text.strip()
        if sn in ("Heading 1", "样式1"):
            in_appendix = bool(appendix_re.match(t_strip))
        if sn.startswith("Heading") or in_appendix or not t_strip:
            continue

        runs = list(p.runs)
        if not runs:
            continue
        chars = []
        for r in runs:
            r_rPr = r._element.find(qn('w:rPr'))
            for ch in (r.text or ""):
                chars.append((ch, r_rPr))
        full_text = "".join(c[0] for c in chars)

        matches = list(_CITE_NUM_RE.finditer(full_text))
        if not matches:
            continue

        # check at least one citation is resolvable
        has_valid = False
        for mat in matches:
            parts = _parse_cite_structure(mat.group(1))
            all_nums = []
            for pt in parts:
                if pt[0] == 'num':
                    all_nums.append(pt[1])
                elif pt[0] == 'range':
                    all_nums.extend(pt[1])
            if all(n in bookmark_map for n in all_nums):
                has_valid = True
                break
        if not has_valid:
            continue

        # clear runs and rebuild
        p_el = p._element
        for child in list(p_el):
            if child.tag != qn('w:pPr'):
                p_el.remove(child)

        pos = 0
        for mat in matches:
            # text before citation
            if mat.start() > pos:
                _append_char_segment(p_el, chars[pos:mat.start()])

            parts = _parse_cite_structure(mat.group(1))
            all_nums = []
            for pt in parts:
                if pt[0] == 'num':
                    all_nums.append(pt[1])
                elif pt[0] == 'range':
                    all_nums.extend(pt[1])
            cite_rPr = chars[mat.start()][1]

            if all(n in bookmark_map for n in all_nums):
                # replace with REF fields
                p_el.append(_make_text_run_el('[', cite_rPr))
                for j, pt in enumerate(parts):
                    if pt[0] == 'sep':
                        p_el.append(_make_text_run_el(',', cite_rPr))
                    elif pt[0] == 'num':
                        bm = bookmark_map[pt[1]]
                        for fel in _make_field_runs(f'REF {bm} \\h', str(pt[1]), cite_rPr):
                            p_el.append(fel)
                    elif pt[0] == 'range':
                        bm_s = bookmark_map[pt[1][0]]
                        bm_e = bookmark_map[pt[1][1]]
                        for fel in _make_field_runs(f'REF {bm_s} \\h', str(pt[1][0]), cite_rPr):
                            p_el.append(fel)
                        p_el.append(_make_text_run_el('-', cite_rPr))
                        for fel in _make_field_runs(f'REF {bm_e} \\h', str(pt[1][1]), cite_rPr):
                            p_el.append(fel)
                p_el.append(_make_text_run_el(']', cite_rPr))
            else:
                # keep as plain text
                _append_char_segment(p_el, chars[mat.start():mat.end()])

            pos = mat.end()

        # remaining text
        if pos < len(chars):
            _append_char_segment(p_el, chars[pos:])


# ---------------------------------------------------------------------------
# Main formatting entry point
# ---------------------------------------------------------------------------

def apply_format(input_path, output_path, config=None, config_path=None):
    if config is None:
        config, config_path = resolve_config(input_path=input_path)
    cfg = config

    latin = cfg["fonts"]["latin"]
    body_font = cfg["fonts"]["body"]
    body_size = Pt(cfg["sizes"]["body"])
    body_ls = cfg["body"]["line_spacing"]
    body_indent = Pt(cfg["body"]["first_line_indent"])
    body_align = _ALIGN_MAP.get(cfg["body"]["align"])  # None when "keep"

    h1_font = cfg["fonts"]["h1"]
    h1_size = Pt(cfg["sizes"]["h1"])
    h2_font = cfg["fonts"]["h2"]
    h2_size = Pt(cfg["sizes"]["h2"])
    h3_font = cfg["fonts"]["h3"]
    h3_size = Pt(cfg["sizes"]["h3"])
    h4_font = cfg["fonts"]["h4"]
    h4_size = Pt(cfg["sizes"]["h4"])

    def _bold_val(v):
        return None if v == "keep" else v

    h1_bold = _bold_val(cfg["headings"]["h1"]["bold"])
    h1_align = _ALIGN_MAP.get(cfg["headings"]["h1"]["align"])
    h2_bold = _bold_val(cfg["headings"]["h2"]["bold"])
    h2_align = _ALIGN_MAP.get(cfg["headings"]["h2"]["align"])
    h3_bold = _bold_val(cfg["headings"]["h3"]["bold"])
    h3_align = _ALIGN_MAP.get(cfg["headings"]["h3"]["align"])
    h4_bold = _bold_val(cfg["headings"]["h4"]["bold"])
    h4_align = _ALIGN_MAP.get(cfg["headings"]["h4"]["align"])

    caption_size = Pt(cfg["sizes"]["caption"])
    note_size = Pt(cfg["sizes"]["note"])
    fn_size = Pt(cfg["sizes"]["footnote"])

    st_map = _get_special_title_map(cfg)
    sec = cfg["sections"]
    ref_key = "参考文献"
    toc_key = "目录"

    doc = Document(input_path)

    # Auto-detect headings in Normal-style paragraphs
    auto_changes = auto_assign_heading_styles(doc, cfg)
    if auto_changes:
        print(f"自动识别标题 ({len(auto_changes)} 个):", file=sys.stderr)
        for c in auto_changes:
            print(c, file=sys.stderr)

    try:
        validate_structure(doc, cfg)
    except Exception as exc:
        print(f"结构检查出错（已跳过，继续排版）: {exc}", file=sys.stderr)
    normalize_sections(doc, cfg)

    # Auto-renumber headings if enabled (skip when auto-assign detected headings,
    # because auto-assign preserves original numbering from the document)
    renum_changes = []
    if sec.get("renumber_headings", False) and not auto_changes:
        renum_changes = renumber_headings(doc, cfg)

    # Normalize heading number-to-title spacing (1字距)
    normalize_heading_spacing(doc, cfg)

    # Base styles
    for style_name in ["Normal", "Body Text", "First Paragraph", "_Style 2"]:
        if style_name in doc.styles:
            set_style_font(doc.styles[style_name], east_asia=body_font,
                           size_pt=body_size, bold=False, latin=latin)

    def _set_heading_style(style_name, font, size, bold, align, hcfg):
        if style_name not in doc.styles:
            return
        st = doc.styles[style_name]
        set_style_font(st, east_asia=font, size_pt=size, bold=bold, latin=latin)
        if align is not None:
            st.paragraph_format.alignment = align
        sb = hcfg.get("space_before", 0)
        sa = hcfg.get("space_after", 0)
        if sb >= 0:
            st.paragraph_format.space_before = Pt(sb * 12)
        if sa >= 0:
            st.paragraph_format.space_after = Pt(sa * 12)

    _set_heading_style("Heading 1", h1_font, h1_size, h1_bold, h1_align, cfg["headings"]["h1"])
    _set_heading_style("Heading 2", h2_font, h2_size, h2_bold, h2_align, cfg["headings"]["h2"])
    _set_heading_style("Heading 3", h3_font, h3_size, h3_bold, h3_align, cfg["headings"]["h3"])
    _set_heading_style("Heading 4", h4_font, h4_size, h4_bold, h4_align, cfg["headings"]["h4"])

    if "TOC Heading" in doc.styles:
        st = doc.styles["TOC Heading"]
        toc_h_font = cfg["toc"].get("h1_font", h1_font)
        toc_h_size = Pt(cfg["toc"].get("h1_font_size", cfg["sizes"]["h1"]))
        set_style_font(st, east_asia=toc_h_font, size_pt=toc_h_size, bold=True, latin=latin)
        st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _toc_h_sb = cfg["toc"].get("space_before", 0)
        _toc_h_sa = cfg["toc"].get("space_after", 0)
        st.paragraph_format.space_before = Pt(_toc_h_sb * 12)
        st.paragraph_format.space_after = Pt(_toc_h_sa * 12)

    ensure_toc_styles(doc, cfg)

    toc_content_font = cfg["toc"].get("font", body_font)
    toc_content_size = Pt(cfg["toc"].get("font_size", cfg["sizes"]["body"]))
    toc_h1_font = cfg["toc"].get("h1_font", cfg["fonts"]["h1"])
    toc_h1_size = Pt(cfg["toc"].get("h1_font_size", cfg["sizes"]["h1"]))
    toc_content_ls = cfg["toc"].get("line_spacing", body_ls)
    toc_sb = Pt(cfg["toc"].get("space_before", 0) * 12)
    toc_sa = Pt(cfg["toc"].get("space_after", 0) * 12)

    for para in doc.paragraphs:
        sn = para.style.name if para.style else ""
        if sn.lower().startswith("toc ") or sn == "样式3":
            is_toc1 = sn.lower() == "toc 1"
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.line_spacing = toc_content_ls
            para.paragraph_format.space_before = toc_sb
            para.paragraph_format.space_after = toc_sa
            ea = toc_h1_font if is_toc1 else toc_content_font
            sz = toc_h1_size if is_toc1 else toc_content_size
            set_para_runs_font(para, east_asia=ea, size_pt=sz,
                               bold=False, latin=latin)

    for name in ["Footnote Text", "Footnote Reference"]:
        if name in doc.styles:
            set_style_font(doc.styles[name], east_asia=body_font, size_pt=fn_size,
                           bold=False, latin=latin)
    if "Footnote Text" in doc.styles:
        ft = doc.styles["Footnote Text"]
        ft.paragraph_format.line_spacing = cfg["footnote"]["line_spacing"]
        ft.paragraph_format.first_line_indent = Pt(0)
        _fn_align = _ALIGN_MAP.get(cfg["footnote"].get("align", "justify"))
        if _fn_align is not None:
            ft.paragraph_format.alignment = _fn_align

    for name in ["Hyperlink", "超链接"]:
        if name in [s.name for s in doc.styles]:
            st = doc.styles[name]
            st.font.color.rgb = RGBColor(0, 0, 0)
            st.font.underline = False

    # Global paragraph defaults
    _heading_styles = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "样式1"}
    for para in doc.paragraphs:
        sn = para.style.name if para.style else ""
        if sn.lower().startswith("toc ") or sn in ("样式3", "TOC Heading"):
            continue  # TOC entries already formatted above
        if sn in _heading_styles:
            # Headings: only zero spacing if config doesn't say "keep original" (-1)
            hkey = {"Heading 1": "h1", "样式1": "h1", "Heading 2": "h2",
                    "Heading 3": "h3", "Heading 4": "h4"}.get(sn, "h1")
            hcfg = cfg["headings"].get(hkey, {})
            if hcfg.get("space_before", 0) >= 0:
                para.paragraph_format.space_before = Pt(0)
            if hcfg.get("space_after", 0) >= 0:
                para.paragraph_format.space_after = Pt(0)
        else:
            para.paragraph_format.space_before = Pt(cfg["body"].get("space_before", 0) * 12)
            para.paragraph_format.space_after = Pt(cfg["body"].get("space_after", 0) * 12)
        pf = para.paragraph_format
        if para.style and para.style.name in ["Normal", "Body Text", "First Paragraph", "_Style 2"]:
            if body_align is not None:
                pf.alignment = body_align
            pf.first_line_indent = body_indent
            pf.line_spacing = body_ls
            set_para_runs_font(para, east_asia=body_font, size_pt=body_size,
                               bold=False, latin=latin)

    # Front matter
    fm_mode = cfg.get("front_matter", {}).get("mode", "auto")
    has_fm = (fm_mode == "format") or \
             (fm_mode == "auto" and _detect_front_matter(doc, cfg))

    cn_kw_para = None
    en_kw_para = None

    if has_fm:
        first_h1_idx = None
        for i, para in enumerate(doc.paragraphs):
            if is_heading(para, 1):
                first_h1_idx = i
                break
        if first_h1_idx is None:
            first_h1_idx = len(doc.paragraphs)

        front = doc.paragraphs[:first_h1_idx]
        non_empty = [p for p in front if p.text.strip()]

        if non_empty:
            abstract_display = _find_special_display(cfg, "摘要")
            p = non_empty[0]
            p.text = abstract_display
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = body_ls
            set_para_runs_font(p, east_asia=h1_font, size_pt=h1_size, bold=True, latin=latin)

        cn_kw_re = sec.get("cn_keywords_pattern", r"^\s*关键词\s*[：:]")
        en_abs_re = sec.get("en_abstract_pattern", r"(?i)^\s*Abstract\s*[：:]")
        en_kw_re = sec.get("en_keywords_pattern", r"(?i)^\s*Key\s*words\s*[：:]")

        past_abstract = False
        en_title_seen = False

        for p in non_empty[1:]:
            t = p.text.strip()
            if t.startswith("关键词"):
                cn_kw_para = p
                normalized = normalize_cn_keywords(t) or t
                content = normalized.split("：", 1)[1] if "：" in normalized else ""
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = body_ls
                r1 = p.add_run("关键词：")
                set_run_font(r1, east_asia=h1_font, size_pt=body_size, bold=True, latin=latin)
                r2 = p.add_run(content)
                set_run_font(r2, east_asia=body_font, size_pt=body_size, bold=False, latin=latin)
            elif re.match(r"^\s*Abstract\s*:", t, flags=re.I):
                past_abstract = True
                content = re.sub(r"^\s*Abstract\s*:\s*", "", t, flags=re.I)
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = body_ls
                r1 = p.add_run("Abstract: ")
                set_run_font(r1, east_asia=latin, size_pt=body_size, bold=True, latin=latin)
                r2 = p.add_run(content)
                set_run_font(r2, east_asia=latin, size_pt=body_size, bold=False, latin=latin)
            elif re.match(r"^\s*Key\s*words\s*:", t, flags=re.I):
                en_kw_para = p
                normalized = normalize_en_keywords(t) or t
                content = re.sub(r"^\s*Key\s*words\s*:\s*", "", normalized, flags=re.I)
                p.clear()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = body_ls
                r1 = p.add_run("Key words: ")
                set_run_font(r1, east_asia=latin, size_pt=body_size, bold=True, latin=latin)
                r2 = p.add_run(content)
                set_run_font(r2, east_asia=latin, size_pt=body_size, bold=False, latin=latin)
            elif not past_abstract and not contains_cjk(t) and not re.match(r"^\s*(Abstract|Key\s*words)\s*:", t, re.I) and len(t) > 20 and not re.match(r"^[\(（]", t):
                en_title_seen = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = body_ls
                set_para_runs_font(p, east_asia=latin, size_pt=h1_size, bold=True, latin=latin)
            elif not past_abstract and en_title_seen and not contains_cjk(t) and not re.match(r"^[\(（]", t) and not re.match(r"^\s*(Abstract|Key\s*words)\s*:", t, re.I):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = body_ls
                set_para_runs_font(p, east_asia=latin, size_pt=body_size, bold=False, latin=latin)
            elif not past_abstract and re.match(r"^[\(（]", t) and re.search(r"(China|University|College)", t, re.I):
                new_t = t
                if new_t.startswith("("):
                    new_t = "（" + new_t[1:]
                if new_t.endswith(")"):
                    new_t = new_t[:-1] + "）"
                if new_t != t:
                    p.text = new_t
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = body_ls
                set_para_runs_font(p, east_asia=latin, size_pt=body_size, bold=False, latin=latin)
            else:
                if contains_cjk(t):
                    if body_align is not None:
                        p.alignment = body_align
                    p.paragraph_format.first_line_indent = body_indent
                    p.paragraph_format.line_spacing = body_ls
                    set_para_runs_font(p, east_asia=body_font, size_pt=body_size,
                                       bold=False, latin=latin)
                elif t:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.line_spacing = body_ls
                    set_para_runs_font(p, east_asia=latin, size_pt=body_size,
                                       bold=False, latin=latin)

        if cn_kw_para is not None:
            insert_page_break_after(cn_kw_para)
        if en_kw_para is not None:
            insert_page_break_after(en_kw_para)

    def _apply_heading_para(para, align, hcfg, font, size, bold):
        if align is not None:
            para.alignment = align
        para.paragraph_format.first_line_indent = Pt(0)
        para.paragraph_format.line_spacing = body_ls
        sb = hcfg.get("space_before", 0)
        sa = hcfg.get("space_after", 0)
        if sb >= 0:
            para.paragraph_format.space_before = Pt(sb * 12)
        if sa >= 0:
            para.paragraph_format.space_after = Pt(sa * 12)
        set_para_runs_font(para, east_asia=font, size_pt=size, bold=bold, latin=latin)

    # Heading-level direct formatting (config-driven special title mapping)
    for para in doc.paragraphs:
        t = para.text.strip()
        t_nospace = t.replace(" ", "").replace("　", "")
        sn = para.style.name if para.style else ""

        # Non-heading TOC title
        if t_nospace in st_map and sn not in ("Heading 1", "Heading 2", "Heading 3", "Heading 4"):
            entry = st_map[t_nospace]
            para.text = entry["display"]
            para.alignment = _ALIGN_MAP.get(entry.get("align", "center"), WD_ALIGN_PARAGRAPH.CENTER)
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.line_spacing = body_ls
            _h1cfg = cfg["headings"]["h1"]
            _sb = _h1cfg.get("space_before", 0)
            _sa = _h1cfg.get("space_after", 0)
            if _sb >= 0:
                para.paragraph_format.space_before = Pt(_sb * 12)
            if _sa >= 0:
                para.paragraph_format.space_after = Pt(_sa * 12)
            set_para_runs_font(para, east_asia=h1_font, size_pt=h1_size,
                               bold=True, latin=latin)
        elif sn == "Heading 1":
            _apply_heading_para(para, h1_align, cfg["headings"]["h1"], h1_font, h1_size, h1_bold)
            if t_nospace in st_map:
                entry = st_map[t_nospace]
                para.text = entry["display"]
                para.alignment = _ALIGN_MAP.get(entry.get("align", "center"),
                                                WD_ALIGN_PARAGRAPH.CENTER)
                set_para_runs_font(para, east_asia=h1_font, size_pt=h1_size,
                                   bold=True, latin=latin)
            elif t.startswith("附录"):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_para_runs_font(para, east_asia=h1_font, size_pt=h1_size,
                                   bold=True, latin=latin)
        elif sn == "Heading 2":
            _apply_heading_para(para, h2_align, cfg["headings"]["h2"], h2_font, h2_size, h2_bold)
        elif sn == "Heading 3":
            _apply_heading_para(para, h3_align, cfg["headings"]["h3"], h3_font, h3_size, h3_bold)
        elif sn == "Heading 4":
            _apply_heading_para(para, h4_align, cfg["headings"]["h4"], h4_font, h4_size, h4_bold)
        elif sn == "样式1":
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.line_spacing = body_ls
            sb = cfg["headings"]["h1"].get("space_before", 0)
            sa = cfg["headings"]["h1"].get("space_after", 0)
            if sb >= 0:
                para.paragraph_format.space_before = Pt(sb * 12)
            if sa >= 0:
                para.paragraph_format.space_after = Pt(sa * 12)
            if t_nospace in st_map:
                entry = st_map[t_nospace]
                para.text = entry["display"]
                para.alignment = _ALIGN_MAP.get(entry.get("align", "center"),
                                                WD_ALIGN_PARAGRAPH.CENTER)
                set_para_runs_font(para, east_asia=h1_font, size_pt=h1_size,
                                   bold=True, latin=latin)
            elif t.startswith("附录"):
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_para_runs_font(para, east_asia=h1_font, size_pt=h1_size,
                                   bold=True, latin=latin)

    # Reference list
    ref_cfg = cfg["references"]
    in_refs = False
    for para in doc.paragraphs:
        t = para.text.strip().replace(" ", "").replace("　", "")
        sn = para.style.name if para.style else ""
        is_h1_like = sn == "Heading 1" or sn == "样式1"
        if is_h1_like and t == ref_key.replace(" ", "").replace("\u3000", ""):
            in_refs = True
            continue
        if is_h1_like and in_refs:
            in_refs = False
        if in_refs and para.text.strip() and not (para.style and para.style.name.startswith("Heading")):
            para.paragraph_format.first_line_indent = Pt(ref_cfg["first_line_indent"])
            para.paragraph_format.left_indent = Pt(ref_cfg["left_indent"])
            para.paragraph_format.line_spacing = body_ls
            set_para_runs_font(para, east_asia=body_font, size_pt=body_size,
                               bold=False, latin=latin)

    # Figure/table captions
    cap_cfg = cfg.get("captions", {})
    fig_pat = cap_cfg.get("figure_pattern", r"^图\s*\d")
    tbl_pat = cap_cfg.get("table_pattern", r"^(续)?表\s*\d")
    subfig_pat = cap_cfg.get("subfigure_pattern", r"^\([a-z]\)")
    note_pat = cap_cfg.get("note_pattern", r"^注[：:]")

    # Auto-fix caption spacing: "图1xxx" → "图1 xxx", "表2yyy" → "表2 yyy"
    _cap_space_re = re.compile(r"^((?:图|表|Figure|Table)\s*[A-Z]?\d+)(\S)", re.I)
    for para in doc.paragraphs:
        t = para.text.strip()
        m = _cap_space_re.match(t)
        if m:
            para.text = m.group(1) + " " + t[m.end(1):]

    spacing_line = Pt(cfg["sizes"]["body"])
    source_pat = r"^(资料)?来源\s*[：:]"

    for para in doc.paragraphs:
        t = para.text.strip()
        if re.match(fig_pat, t) or re.match(r"^Figure\s*\d", t, re.I) or re.match(r"^图[A-Z]\d+", t):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.line_spacing = body_ls
            para.paragraph_format.space_after = spacing_line
            set_para_runs_font(para, east_asia=body_font, size_pt=caption_size,
                               bold=False, latin=latin)
        elif re.match(tbl_pat, t) or re.match(r"^Table\s*\d", t, re.I) or re.match(r"^(续)?表[A-Z]\d+", t):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.line_spacing = body_ls
            para.paragraph_format.space_before = spacing_line
            set_para_runs_font(para, east_asia=body_font, size_pt=caption_size,
                               bold=False, latin=latin)
        elif re.match(subfig_pat, t):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.line_spacing = body_ls
            set_para_runs_font(para, east_asia=body_font, size_pt=caption_size,
                               bold=False, latin=latin)
        elif re.match(note_pat, t) or re.match(source_pat, t):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.line_spacing = body_ls
            set_para_runs_font(para, east_asia=body_font, size_pt=note_size,
                               bold=False, latin=latin)

    # Keep captions with their figures/tables + spacing around figures/tables
    if cap_cfg.get("keep_with_next", True):
        body_el = doc.element.body
        children = list(body_el)
        for i, el in enumerate(children):
            if el.tag == qn("w:tbl"):
                if i > 0 and children[i - 1].tag == qn("w:p"):
                    prev_text = "".join(
                        (nd.text or "") for nd in children[i - 1].iter(qn("w:t"))
                    ).strip()
                    if re.match(tbl_pat, prev_text) or re.match(r"^Table\s*\d", prev_text, re.I) or re.match(r"^(续)?表[A-Z]\d+", prev_text):
                        _ensure_keep_next(children[i - 1])
                # After table: add space_before on next text paragraph
                if i + 1 < len(children) and children[i + 1].tag == qn("w:p"):
                    nt = "".join(
                        (nd.text or "") for nd in children[i + 1].iter(qn("w:t"))
                    ).strip()
                    if nt and not re.match(note_pat, nt) and not re.match(source_pat, nt):
                        _set_para_spacing(children[i + 1], "before", spacing_line)
            elif el.tag == qn("w:p") and el.findall(".//" + qn("w:drawing")):
                if i + 1 < len(children) and children[i + 1].tag == qn("w:p"):
                    next_text = "".join(
                        (nd.text or "") for nd in children[i + 1].iter(qn("w:t"))
                    ).strip()
                    if re.match(fig_pat, next_text) or re.match(r"^Figure\s*\d", next_text, re.I) or re.match(r"^图[A-Z]\d+", next_text):
                        _ensure_keep_next(el)
                # Before figure: add space_before on the figure paragraph
                if i > 0 and children[i - 1].tag == qn("w:p"):
                    pt = "".join(
                        (nd.text or "") for nd in children[i - 1].iter(qn("w:t"))
                    ).strip()
                    if pt and not re.match(fig_pat, pt) and not re.match(subfig_pat, pt):
                        _set_para_spacing(el, "before", spacing_line)

    # Validate caption numbering
    warnings = []
    if cap_cfg.get("check_numbering", True):
        warnings = _check_caption_numbering(doc, fig_pat, tbl_pat, cfg)

    # Citation / reference cross-check
    try:
        warnings.extend(check_citations(doc, cfg))
    except Exception as exc:
        print(f"引用检查出错（已跳过）: {exc}", file=sys.stderr)

    # Fix citation comma spacing: "(Author,YEAR)" → "(Author, YEAR)"
    _cite_comma = re.compile(r",\s*((?:19|20)\d{2})")
    for para in doc.paragraphs:
        for run in para.runs:
            old = run.text
            new = _cite_comma.sub(r", \1", old)
            if new != old:
                run.text = new
                print(f"  引用逗号修正: \"{old.strip()[:40]}\" → \"{new.strip()[:40]}\"")

    # Cross-reference fields for numbered citations
    try:
        _apply_ref_crosslinks(doc, cfg)
    except Exception as exc:
        print(f"交叉引用创建出错（已跳过）: {exc}", file=sys.stderr)

    # Table formatting
    tbl_cfg = cfg["table"]
    tbl_cell_align = _ALIGN_MAP.get(tbl_cfg.get("cell_align", "center"))
    for table in doc.tables:
        # Autofit: table width = 100% of page
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is None:
            tblLayout = OxmlElement('w:tblLayout')
            tblPr.append(tblLayout)
        tblLayout.set(qn('w:type'), 'autofit')

        rows = len(table.rows)
        for r_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for p in cell.paragraphs:
                    if tbl_cell_align is not None:
                        p.alignment = tbl_cell_align
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.first_line_indent = Pt(0)
                    p.paragraph_format.line_spacing = tbl_cfg["line_spacing"]
                    set_para_runs_font(p, east_asia=body_font, size_pt=caption_size,
                                       bold=False, latin=latin)

                clear_table_border(cell, "left")
                clear_table_border(cell, "right")
                clear_table_border(cell, "insideV")
                clear_table_border(cell, "insideH")

                if r_idx == 0:
                    set_table_border(cell, "top", sz=tbl_cfg["top_border_sz"])
                    set_table_border(cell, "bottom", sz=tbl_cfg["header_border_sz"])
                if r_idx == rows - 1:
                    set_table_border(cell, "bottom", sz=tbl_cfg["bottom_border_sz"])

    # TOC
    insert_toc(doc, cfg)

    # Page breaks before body H1s
    toc_match = _find_special_display(cfg, "目录", raw=True)
    first_body_h1_seen = False
    for para in doc.paragraphs:
        sn = para.style.name if para.style else ""
        if sn in ("Heading 1", "样式1"):
            t = para.text.strip().replace(" ", "").replace("\u3000", "")
            if t == toc_match:
                continue
            if not first_body_h1_seen:
                first_body_h1_seen = True
                continue
            para.paragraph_format.page_break_before = True

    # Page numbering
    setup_page_numbers(doc, cfg)
    try:
        setup_headers(doc, cfg)
    except Exception as e:
        print(f"  [警告] 页眉设置出错，已跳过: {e}", file=sys.stderr)

    # Cover page
    custom_cover = cfg.get("cover", {}).get("custom_docx", "")
    if custom_cover and os.path.isfile(custom_cover):
        insert_cover_and_declaration(doc, cfg, config_path, skip_cover=True)
        try:
            insert_custom_cover(doc, custom_cover)
        except Exception as exc:
            print(f"自定义封面插入出错（已跳过）: {exc}", file=sys.stderr)
    elif cfg["cover"]["enabled"] and not _has_cover(doc, cfg):
        insert_cover_and_declaration(doc, cfg, config_path)

    doc.save(output_path)
    patch_theme_fonts(output_path, cfg)
    if renum_changes:
        warnings.append("标题编号已自动修正:")
        warnings.extend(renum_changes)
    return warnings


def patch_theme_fonts(docx_path, cfg):
    import xml.etree.ElementTree as ET
    theme = cfg.get("theme_fonts", {})
    theme_latin = theme.get("latin", "Times New Roman")
    theme_hans = theme.get("hans", "宋体")

    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    ET.register_namespace("a", a_ns)
    ns = {"a": a_ns}

    fd, tmp_path = tempfile.mkstemp(suffix=".docx", dir=os.path.dirname(docx_path))
    os.close(fd)
    try:
        with zipfile.ZipFile(docx_path, "r") as zin, \
             zipfile.ZipFile(tmp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/theme/theme1.xml":
                    root = ET.fromstring(data)
                    for minor in root.findall(".//a:minorFont", ns):
                        lat = minor.find("a:latin", ns)
                        if lat is not None:
                            lat.set("typeface", theme_latin)
                        for font in minor.findall("a:font", ns):
                            if font.get("script") == "Hans":
                                font.set("typeface", theme_hans)
                    for major in root.findall(".//a:majorFont", ns):
                        lat = major.find("a:latin", ns)
                        if lat is not None:
                            lat.set("typeface", theme_latin)
                        for font in major.findall("a:font", ns):
                            if font.get("script") == "Hans":
                                font.set("typeface", theme_hans)
                    data = ET.tostring(root, encoding="unicode").encode("utf-8")
                zout.writestr(item, data)
        os.replace(tmp_path, docx_path)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Universal thesis formatter")
    parser.add_argument("--input", required=True, help="Input docx")
    parser.add_argument("--output", required=True, help="Output docx")
    parser.add_argument("--config", help="Path to thesis_config.yaml")
    args = parser.parse_args()

    cfg, cfg_path = resolve_config(cli_config=args.config, input_path=args.input)
    apply_format(args.input, args.output, config=cfg, config_path=cfg_path)
    print(f"OK {args.output}")


if __name__ == "__main__":
    main()
