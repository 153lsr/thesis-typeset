import copy
import shutil
import tempfile
import unittest
import warnings
from pathlib import Path
from types import SimpleNamespace
from unittest import mock

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

import thesis_format_cli
import word_postprocess
from thesis_config import DEFAULT_CONFIG
from thesis_formatter import formatter, headers, headings, page
from thesis_formatter.formatter import apply_format


class _FakeStdin:
    def __init__(self, is_tty=True, exc=None):
        self._is_tty = is_tty
        self._exc = exc

    def isatty(self):
        if self._exc is not None:
            raise self._exc
        return self._is_tty


def _pg_num_start(section):
    pg_num = section._sectPr.find(qn("w:pgNumType"))
    return pg_num.get(qn("w:start")) if pg_num is not None else None


class RuntimeSafetyTests(unittest.TestCase):
    def test_should_prompt_before_exit_only_for_interactive_frozen_exe(self):
        with mock.patch.object(thesis_format_cli.sys, "stdin", _FakeStdin(True)):
            with mock.patch.object(thesis_format_cli.sys, "frozen", True, create=True):
                self.assertTrue(thesis_format_cli.should_prompt_before_exit())

            with mock.patch.object(thesis_format_cli.sys, "frozen", False, create=True):
                self.assertFalse(thesis_format_cli.should_prompt_before_exit())

        with mock.patch.object(thesis_format_cli.sys, "stdin", _FakeStdin(exc=RuntimeError("boom"))):
            with mock.patch.object(thesis_format_cli.sys, "frozen", True, create=True):
                self.assertFalse(thesis_format_cli.should_prompt_before_exit())

    def test_terminate_process_targets_specific_pid(self):
        with mock.patch.object(word_postprocess.subprocess, "run", return_value=SimpleNamespace(returncode=0)) as run_mock:
            self.assertTrue(word_postprocess._terminate_process(4321))

        run_mock.assert_called_once_with(
            ["taskkill", "/F", "/PID", "4321"],
            capture_output=True,
            text=True,
            timeout=5,
        )

    def test_normalize_heading_spacing_handles_english_chapter_without_warning(self):
        doc = Document()
        para = doc.add_paragraph("Chapter 2 Introduction", style="Heading 1")

        with warnings.catch_warnings(record=True) as caught:
            warnings.simplefilter("always")
            headings.normalize_heading_spacing(doc, copy.deepcopy(DEFAULT_CONFIG))

        self.assertEqual(para.text, "Chapter 2  Introduction")
        self.assertFalse(
            any(issubclass(item.category, DeprecationWarning) for item in caught),
            caught,
        )

    def test_renumber_headings_accepts_skip_para_ids(self):
        doc = Document()
        preserved = doc.add_paragraph("第9章 保留", style="Heading 1")
        body = doc.add_paragraph("第7章 正文", style="Heading 1")

        changes = headings.renumber_headings(
            doc,
            copy.deepcopy(DEFAULT_CONFIG),
            skip_para_ids={id(preserved._element)},
        )

        self.assertEqual(preserved.text, "第9章 保留")
        self.assertEqual(body.text, "第1章 正文")
        self.assertTrue(any("第7章 正文" in change for change in changes), changes)

    def test_build_insert_cover_vbs_uses_section_break_and_page_setup_copy(self):
        vbs = formatter._build_insert_cover_vbs()

        self.assertIn("Const wdSectionBreakNextPage = 2", vbs)
        self.assertIn("Const wdFormatXMLDocument = 12", vbs)
        self.assertIn("fso.CopyFile targetPath, tempBodyPath, True", vbs)
        self.assertIn("mergedDoc.SaveAs2 targetPath, wdFormatXMLDocument", vbs)
        self.assertIn("objWord.Selection.InsertBreak wdSectionBreakNextPage", vbs)
        self.assertIn("objWord.Selection.InsertFile tempBodyPath", vbs)
        self.assertIn("ClearSectionHeaderFooter mergedDoc.Sections(1)", vbs)

    def test_setup_page_numbers_skips_custom_cover_section(self):
        doc = Document()
        doc.add_paragraph("自定义封面")
        doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_paragraph("摘要")
        doc.add_paragraph("第1章 绪论", style="Heading 1")

        cfg = copy.deepcopy(DEFAULT_CONFIG)
        cfg.setdefault("_runtime", {})["custom_cover_sections"] = 1

        page.setup_page_numbers(doc, cfg)

        self.assertGreaterEqual(len(doc.sections), 3)
        self.assertTrue(all(not p.text for p in doc.sections[0].footer.paragraphs))
        self.assertTrue(any(p.text for p in doc.sections[1].footer.paragraphs))
        self.assertEqual(_pg_num_start(doc.sections[1]), str(cfg["page_numbers"]["front_start"]))
        self.assertEqual(_pg_num_start(doc.sections[-1]), str(cfg["page_numbers"]["body_start"]))

    def test_setup_headers_skips_only_cover_when_scope_is_all(self):
        doc = Document()
        doc.add_paragraph("自定义封面")
        doc.add_section(WD_SECTION.NEW_PAGE)
        doc.add_paragraph("摘要")
        doc.add_paragraph("第1章 绪论", style="Heading 1")

        cfg = copy.deepcopy(DEFAULT_CONFIG)
        cfg.setdefault("_runtime", {})["custom_cover_sections"] = 1
        cfg["header_footer"]["enabled"] = True
        cfg["header_footer"]["scope"] = "all"
        cfg["header_footer"]["different_odd_even"] = False

        page.setup_page_numbers(doc, cfg)
        headers.setup_headers(doc, cfg)

        self.assertTrue(all(not p.text for p in doc.sections[0].header.paragraphs))
        self.assertEqual(doc.sections[1].header.paragraphs[0].text, cfg["header_footer"]["odd_page_text"])
        self.assertEqual(doc.sections[-1].header.paragraphs[0].text, cfg["header_footer"]["odd_page_text"])

    def test_apply_format_inserts_custom_cover_before_page_numbers(self):
        tmpdir = Path(tempfile.mkdtemp(prefix="runtime_safety_cover_"))
        try:
            input_path = tmpdir / "input.docx"
            cover_path = tmpdir / "cover.docx"
            output_path = tmpdir / "output.docx"

            doc = Document()
            doc.add_paragraph("摘要")
            doc.add_paragraph("这是摘要内容。")
            doc.add_paragraph("第1章 绪论", style="Heading 1")
            doc.save(input_path)

            cover_doc = Document()
            cover_doc.add_paragraph("本科毕业论文(或设计)")
            cover_doc.save(cover_path)

            cfg = copy.deepcopy(DEFAULT_CONFIG)
            cfg["cover"]["custom_docx"] = str(cover_path)
            cfg["toc"]["enabled"] = False

            call_order = []
            with mock.patch("thesis_formatter.formatter._insert_cover_via_vbs", side_effect=lambda *_: ((call_order.append("insert") or True), "")):
                with mock.patch("thesis_formatter.formatter.setup_page_numbers", side_effect=lambda *_args: call_order.append(("page", cfg["_runtime"].get("custom_cover_sections")))):
                    with mock.patch("thesis_formatter.formatter.setup_headers", side_effect=lambda *_args: call_order.append(("header", cfg["_runtime"].get("custom_cover_sections")))):
                        with mock.patch("thesis_formatter.formatter.patch_theme_fonts"):
                            apply_format(str(input_path), str(output_path), config=cfg)

            self.assertEqual(call_order[0], "insert")
            self.assertEqual(call_order[1], ("page", 1))
            self.assertEqual(call_order[2], ("header", 1))
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)


if __name__ == "__main__":
    unittest.main()
