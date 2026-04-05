import re

from ._common import (
    is_heading, get_paragraph_heading_level, get_heading_style,
    _ALIGN_MAP, set_para_runs_font, set_run_font,
    _ALL_HEADING_NAMES, _HEADING_STYLE_IDS,
)
from ._titles import _find_special_display, _get_special_title_map
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


_CN_NUMERALS = "零一二三四五六七八九十"
_AR_CHAPTER_PATTERN = r"^第\s*\d+\s*章(?:\s|(?=[\u4e00-\u9fff])|$)"
_CN_CHAPTER_PATTERN = r"^第\s*[一二三四五六七八九十百千零两〇]+\s*章(?:\s|(?=[\u4e00-\u9fff])|$)"


def _chapter_patterns(cfg, text_first=False):
    sec = cfg.get("sections", {})
    configured = sec.get("chapter_pattern", _AR_CHAPTER_PATTERN)
    ordered = [_CN_CHAPTER_PATTERN, _AR_CHAPTER_PATTERN, configured] if text_first else [configured, _CN_CHAPTER_PATTERN, _AR_CHAPTER_PATTERN]
    seen = set()
    for pat in ordered:
        if pat and pat not in seen:
            seen.add(pat)
            yield pat


def _match_chapter_heading(text, cfg, text_first=False):
    for pat in _chapter_patterns(cfg, text_first=text_first):
        m = re.match(pat, text)
        if m:
            return m
    return None


def _int_to_cn(n):
    if n <= 10:
        return _CN_NUMERALS[n]
    if n < 20:
        return "十" + (_CN_NUMERALS[n - 10] if n > 10 else "")
    tens, ones = divmod(n, 10)
    return (_CN_NUMERALS[tens] if tens > 1 else "") + "十" + \
           (_CN_NUMERALS[ones] if ones else "")


def _renumber_h1_text(text, new_num, pattern):
    if re.search(r"第\s*(?:\d+|[一二三四五六七八九十百千零两〇]+)\s*章", text):
        return re.sub(r"(第\s*)(?:\d+|[一二三四五六七八九十百千零两〇]+)(\s*章)", fr"\g<1>{new_num}\2", text, count=1)
    if re.search(r"(?i)Chapter\s+\d+", text):
        return re.sub(r"(?i)(Chapter\s+)\d+", fr"\g<1>{new_num}", text, count=1)
    if re.match(r"^[一二三四五六七八九十百]+、", text):
        cn = _int_to_cn(new_num)
        return re.sub(r"^[一二三四五六七八九十百]+", cn, text, count=1)
    if re.match(r"^\d+\s", text):
        return re.sub(r"^\d+", str(new_num), text, count=1)
    return text


def _renumber_sub_text(text, prefix):
    if re.match(r"^（[一二三四五六七八九十百]+）", text):
        return re.sub(r"^（[一二三四五六七八九十百]+）", prefix, text, count=1)
    return re.sub(r"^[\d.]+", prefix, text, count=1)


def _replace_para_text(para, new_text):
    if para.runs:
        first_run = para.runs[0]
        font_props = {
            "name": first_run.font.name,
            "size": first_run.font.size,
            "bold": first_run.font.bold,
        }
        ea = first_run.font.element.find(
            ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
        ea_name = ea.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia") if ea is not None else None
    else:
        font_props = None
        ea_name = None

    para.text = new_text

    if font_props and para.runs:
        r = para.runs[0]
        r.font.name = font_props["name"]
        r.font.size = font_props["size"]
        r.font.bold = font_props["bold"]
        if ea_name:
            rFonts = r.font.element.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                r.font.element.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), ea_name)





def renumber_headings(doc, cfg, skip_para_ids=None):
    skip_para_ids = set(skip_para_ids or ())
    sec = cfg.get("sections", {})
    text_first = bool(cfg.get("toc", {}).get("only_insert", False))
    appendix_pat = re.compile(sec.get("appendix_pattern", r"^附录\s*[A-Z]"))
    h2_pat = re.compile(sec.get("h2_pattern", r"^\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    h3_pat = re.compile(sec.get("h3_pattern", r"^\d+\.\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    h4_pat = re.compile(sec.get("h4_pattern", r"^\d+\.\d+\.\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))

    st_map = _get_special_title_map(cfg)
    special_set = set(st_map.keys())
    special_set.update(s.replace(" ", "").replace("\u3000", "")
                       for s in sec.get("special_h1", []))

    changes = []
    chap_n = 0
    h2_n = h3_n = h4_n = 0
    in_appendix = False

    for para in doc.paragraphs:
        if id(para._element) in skip_para_ids:
            continue
        level = get_paragraph_heading_level(para)
        t = para.text.strip()
        t_nospace = t.replace(" ", "").replace("\u3000", "")
        if not t:
            continue

        if level == 1:
            if t_nospace in special_set:
                continue
            if appendix_pat.match(t):
                in_appendix = True
                continue
            if in_appendix:
                continue
            chapter_match = _match_chapter_heading(t, cfg, text_first=text_first)
            if chapter_match:
                chap_n += 1
                h2_n = h3_n = h4_n = 0
                new_t = _renumber_h1_text(t, chap_n, chapter_match.re.pattern)
                if new_t != t:
                    changes.append(f'  H1: "{t}" → "{new_t}"')
                    _replace_para_text(para, new_t)
        elif level == 2 and not in_appendix:
            if h2_pat.match(t):
                h2_n += 1
                h3_n = h4_n = 0
                new_t = _renumber_sub_text(t, f"{chap_n}.{h2_n}")
                if new_t != t:
                    changes.append(f'  H2: "{t}" → "{new_t}"')
                    _replace_para_text(para, new_t)
        elif level == 3 and not in_appendix:
            if h3_pat.match(t):
                h3_n += 1
                h4_n = 0
                new_t = _renumber_sub_text(t, f"{chap_n}.{h2_n}.{h3_n}")
                if new_t != t:
                    changes.append(f'  H3: "{t}" → "{new_t}"')
                    _replace_para_text(para, new_t)
        elif level == 4 and not in_appendix:
            if h4_pat.match(t):
                h4_n += 1
                new_t = _renumber_sub_text(t, f"{chap_n}.{h2_n}.{h3_n}.{h4_n}")
                if new_t != t:
                    changes.append(f'  H4: "{t}" → "{new_t}"')
                    _replace_para_text(para, new_t)

    return changes



def normalize_heading_spacing(doc, cfg, skip_para_ids=None):
    skip_para_ids = set(skip_para_ids or ())
    text_first = bool(cfg.get("toc", {}).get("only_insert", False))
    st_map = _get_special_title_map(cfg)
    JIJU = "  "

    for para in doc.paragraphs:
        if id(para._element) in skip_para_ids:
            continue
        level = get_paragraph_heading_level(para)
        t = para.text.strip()
        if not t:
            continue
        t_nospace = t.replace(" ", "").replace("\u3000", "")
        chapter_match = _match_chapter_heading(t, cfg, text_first=text_first)
        effective_level = level if level is not None else _infer_heading_level_from_text(t, cfg, text_first=text_first, para=para)

        new_t = None
        if effective_level == 1:
            if t_nospace in st_map:
                continue
            if chapter_match:
                suffix = t[chapter_match.end():].lstrip()
                if suffix:
                    new_t = chapter_match.group(0).rstrip() + JIJU + suffix
            if new_t is None:
                m = re.match(r"(?i)(Chapter\s+\d+)\s*(.*)", t)
                if m and m.group(2):
                    new_t = m.group(1) + JIJU + m.group(2)
            if new_t is None:
                m = re.match(r"([一二三四五六七八九十百]+、)\s*(.*)", t)
                if m and m.group(2):
                    new_t = m.group(1) + JIJU + m.group(2)
        elif effective_level in (2, 3, 4):
            m = re.match(r"(（[一二三四五六七八九十百]+）)\s*(.*)", t)
            if m and m.group(2):
                new_t = m.group(1) + JIJU + m.group(2)
            if new_t is None:
                m = re.match(r"([\d.]+)\s*(.*)", t)
                if m and m.group(2):
                    new_t = m.group(1) + JIJU + m.group(2)
            if new_t is None:
                m = re.match(r"(\(\d+\))\s*(.*)", t)
                if m and m.group(2):
                    new_t = m.group(1) + JIJU + m.group(2)

        if new_t is not None and new_t != t:
            _replace_para_text(para, new_t)


def _paragraph_has_graphics(para):
    p_el = para._element
    for tag in ("w:drawing", "w:pict", "w:object"):
        if p_el.findall(".//" + qn(tag)):
            return True
    return False


def _get_paragraph_outline_level(para):
    def _extract(ppr):
        if ppr is None:
            return None
        outline = ppr.find(qn("w:outlineLvl"))
        if outline is None:
            return None
        val = outline.get(qn("w:val"))
        try:
            level = int(val) + 1
        except (TypeError, ValueError):
            return None
        return level if 1 <= level <= 4 else None

    level = _extract(para._element.find(qn("w:pPr")))
    if level is not None:
        return level

    style = para.style
    checked = set()
    while style is not None:
        style_id = getattr(style, "style_id", None)
        if style_id in checked:
            break
        if style_id:
            checked.add(style_id)
        style_el = getattr(style, "element", None)
        if style_el is not None:
            level = _extract(style_el.find(qn("w:pPr")))
            if level is not None:
                return level
        style = getattr(style, "base_style", None)
    return None







def _clear_paragraph_outline_level(para):
    ppr = para._element.find(qn("w:pPr"))
    if ppr is None:
        return
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is not None:
        ppr.remove(outline)


def _normalize_title_text(text):
    return text.replace(" ", "").replace("\u3000", "")


def _layout_allows_heading_level(para, level):
    if level == 1:
        return True
    text = para.text or ""
    if text[:1] in (" ", "\t", "\u3000"):
        return False
    return para.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT)


def _infer_heading_level_from_text(text, cfg, text_first=False, para=None):
    sec = cfg.get("sections", {})
    appendix_re = re.compile(sec.get("appendix_pattern", r"^附录\s*[A-Z]"))
    h2_re = re.compile(sec.get("h2_pattern", r"^\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    h3_re = re.compile(sec.get("h3_pattern", r"^\d+\.\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    h4_re = re.compile(sec.get("h4_pattern", r"^\d+\.\d+\.\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))

    candidate = None
    if _match_chapter_heading(text, cfg, text_first=text_first):
        candidate = 1
    elif appendix_re.match(text):
        candidate = 1
    elif re.match(r"(?i)^Chapter\s+\d+\b", text):
        candidate = 1
    elif re.match(r"^[一二三四五六七八九十百]+、", text):
        candidate = 1
    elif h4_re.match(text):
        candidate = 4
    elif h3_re.match(text):
        candidate = 3
    elif h2_re.match(text):
        candidate = 2

    if candidate is None:
        return None
    if not _looks_like_auto_heading_text(text, candidate):
        return None
    if para is not None and not _layout_allows_heading_level(para, candidate):
        return None
    return candidate
def _get_rfonts_value(rpr, attr_name):
    if rpr is None:
        return None
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        return None
    return rfonts.get(qn(f"w:{attr_name}"))


def _set_run_rfonts(run, latin=None, east_asia=None):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    if latin:
        rfonts.set(qn("w:ascii"), latin)
        rfonts.set(qn("w:hAnsi"), latin)
    if east_asia:
        rfonts.set(qn("w:eastAsia"), east_asia)


def _materialize_paragraph_look(para):
    style = para.style
    if style is None:
        return

    pf = para.paragraph_format
    spf = style.paragraph_format

    if para.alignment is None and spf.alignment is not None:
        para.alignment = spf.alignment

    for attr in (
        "left_indent",
        "right_indent",
        "first_line_indent",
        "space_before",
        "space_after",
        "line_spacing",
        "keep_together",
        "keep_with_next",
        "page_break_before",
        "widow_control",
    ):
        current = getattr(pf, attr)
        if current is not None:
            continue
        style_value = getattr(spf, attr)
        if style_value is not None:
            setattr(pf, attr, style_value)

    style_font = style.font
    style_rpr = style.element.find(qn("w:rPr"))
    style_ea = _get_rfonts_value(style_rpr, "eastAsia")
    style_latin = _get_rfonts_value(style_rpr, "ascii") or style_font.name
    style_size = style_font.size
    style_bold = style_font.bold
    style_italic = style_font.italic
    style_underline = style_font.underline
    style_color = style_font.color.rgb

    for run in para.runs:
        font = run.font
        run_rpr = run._element.find(qn("w:rPr"))
        latin = _get_rfonts_value(run_rpr, "ascii") or font.name or style_latin
        east_asia = _get_rfonts_value(run_rpr, "eastAsia") or style_ea
        size = font.size or style_size
        bold = style_bold if font.bold is None else font.bold
        italic = style_italic if font.italic is None else font.italic
        underline = style_underline if font.underline is None else font.underline
        color = font.color.rgb or style_color

        if latin is not None:
            font.name = latin
        if size is not None:
            font.size = size
        if bold is not None:
            font.bold = bold
        if italic is not None:
            font.italic = italic
        if underline is not None:
            font.underline = underline
        if color is not None:
            font.color.rgb = color
        if latin or east_asia:
            _set_run_rfonts(run, latin=latin, east_asia=east_asia)


def _resolve_demote_target_style(doc):
    for name in ("Normal", "Body Text", "正文"):
        try:
            style = doc.styles[name]
            if style is not None:
                return style
        except KeyError:
            continue
    for style in doc.styles:
        if style.type == 1 and getattr(style, "style_id", "") not in _HEADING_STYLE_IDS.values():
            return style
    return None


def _is_abstract_heading(text, cfg):
    normalized = _normalize_title_text(text)
    cn_titles = {
        _normalize_title_text("摘要"),
        _normalize_title_text(_find_special_display(cfg, "摘要", raw=True)),
        _normalize_title_text(_find_special_display(cfg, "摘要")),
    }
    return normalized in cn_titles or normalized.lower() == "abstract"


def _is_caption_like_heading(text, cfg):
    cap_cfg = cfg.get("captions", {})
    fig_pat = cap_cfg.get("figure_pattern", r"^图\s*\d")
    tbl_pat = cap_cfg.get("table_pattern", r"^(续)?表\s*\d")
    subfig_pat = cap_cfg.get("subfigure_pattern", r"^\([a-z]\)")
    note_pat = cap_cfg.get("note_pattern", r"^注[：:]")
    source_pat = r"^(资料)?来源\s*[：:]"

    return (
        re.match(fig_pat, text)
        or re.match(r"^Figure\s*\d", text, re.I)
        or re.match(r"^图[A-Z]\d+", text)
        or re.match(tbl_pat, text)
        or re.match(r"^Table\s*\d", text, re.I)
        or re.match(r"^(续)?表[A-Z]\d+", text)
        or re.match(subfig_pat, text)
        or re.match(note_pat, text)
        or re.match(source_pat, text)
    )





def _matches_structured_heading(level, text, cfg):
    sec = cfg.get("sections", {})
    text_first = bool(cfg.get("toc", {}).get("only_insert", False))
    appendix_re = re.compile(sec.get("appendix_pattern", r"^附录\s*[A-Z]"))
    h2_re = re.compile(sec.get("h2_pattern", r"^\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    h3_re = re.compile(sec.get("h3_pattern", r"^\d+\.\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    h4_re = re.compile(sec.get("h4_pattern", r"^\d+\.\d+\.\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    st_map = _get_special_title_map(cfg)
    special_h1_set = set(st_map.keys())
    special_h1_set.update(s.replace(" ", "").replace("\u3000", "") for s in sec.get("special_h1", []))

    normalized = _normalize_title_text(text)
    if level == 1:
        if normalized in special_h1_set:
            return True
        return bool(
            _match_chapter_heading(text, cfg, text_first=text_first)
            or appendix_re.match(text)
            or re.match(r"(?i)^Chapter\s+\d+\b", text)
            or re.match(r"^[一二三四五六七八九十百]+、", text)
        )
    if level == 2:
        return bool(h2_re.match(text) or re.match(r"^（[一二三四五六七八九十百]+）", text) or re.match(r"^\(\d+\)", text))
    if level == 3:
        return bool(h3_re.match(text) or re.match(r"^（[一二三四五六七八九十百]+）", text) or re.match(r"^\(\d+\)", text))
    if level == 4:
        return bool(h4_re.match(text) or re.match(r"^（[一二三四五六七八九十百]+）", text) or re.match(r"^\(\d+\)", text))
    return False


def _looks_like_body_paragraph(text):
    stripped = text.strip()
    if not stripped:
        return False
    if len(stripped) >= 90:
        return True
    if stripped[-1] in "。！？；.!?;":
        return True
    punct_count = sum(stripped.count(ch) for ch in "，,：:；;。！？.!?")
    if len(stripped) >= 45 and punct_count >= 1:
        return True
    if len(stripped) >= 35 and punct_count >= 2:
        return True
    return False


def _needs_body_heading_demote(level, text, cfg, aggressive=False):
    if _matches_structured_heading(level, text, cfg):
        return False
    if aggressive and level in (1, 2, 3, 4):
        return True
    return _looks_like_body_paragraph(text)





def demote_abstract_heading_styles(doc, cfg, include_abstract=True, aggressive_body_demote=False):
    target_style = _resolve_demote_target_style(doc)
    if target_style is None:
        return []

    changes = []
    for para in doc.paragraphs:
        level = get_paragraph_heading_level(para)
        outline_level = _get_paragraph_outline_level(para) if level is None else None
        if level is None and outline_level is None:
            continue
        text = para.text.strip()
        effective_level = level if level is not None else outline_level

        reason = None
        if _paragraph_has_graphics(para):
            reason = "误套标题的图片段落"
        elif not text:
            reason = "误套标题的空白段落"
        elif include_abstract and effective_level == 1 and _is_abstract_heading(text, cfg):
            reason = "摘要标题"
        elif _is_caption_like_heading(text, cfg):
            reason = "误套标题的图表题注"
        elif effective_level in (2, 3, 4) and not _layout_allows_heading_level(para, effective_level):
            reason = "误套标题的正文"
        elif (
            (level is not None or (aggressive_body_demote and outline_level is not None))
            and _needs_body_heading_demote(effective_level, text, cfg, aggressive=aggressive_body_demote)
        ):
            reason = "误套标题的正文"

        if reason is None:
            continue

        _materialize_paragraph_look(para)
        para.style = target_style
        _clear_paragraph_outline_level(para)
        display_text = text or "<空白段落>"
        changes.append(f'  解除{reason} Heading: "{display_text}"')

    return changes


_SENTENCE_ENDINGS = set("。！？；")
_YEAR_LIKE_HEADING_RE = re.compile(r"^\d{4}\s*年")


def _looks_like_auto_heading_text(text, level):
    text = (text or "").strip()
    if not text:
        return False
    if len(text) > 50:
        return False
    if text[-1] in _SENTENCE_ENDINGS:
        return False
    if level == 1 and _YEAR_LIKE_HEADING_RE.match(text):
        return False
    return True

def auto_assign_heading_styles(doc, cfg, skip_para_ids=None, preserve_look=False):
    skip_para_ids = set(skip_para_ids or ())
    sec = cfg.get("sections", {})
    text_first = bool(cfg.get("toc", {}).get("only_insert", False))
    appendix_re = re.compile(sec.get("appendix_pattern", r"^附录\s*[A-Z]"))
    h2_re = re.compile(sec.get("h2_pattern", r"^\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    h3_re = re.compile(sec.get("h3_pattern", r"^\d+\.\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))
    h4_re = re.compile(sec.get("h4_pattern", r"^\d+\.\d+\.\d+\.\d+(\s|(?=[\u4e00-\u9fff]))"))

    st_map = _get_special_title_map(cfg)
    special_h1_set = set()
    for s in sec.get("special_h1", []):
        special_h1_set.add(s.replace(" ", "").replace("\u3000", ""))
    special_h1_set.update(st_map.keys())
    non_outline_front_titles = {
        "摘要",
        _find_special_display(cfg, "摘要").replace(" ", "").replace("\u3000", ""),
        "Abstract",
    }

    changes = []
    for para in doc.paragraphs:
        if id(para._element) in skip_para_ids:
            continue
        level = get_paragraph_heading_level(para)
        t = para.text.strip()
        if not t:
            continue
        if _paragraph_has_graphics(para):
            continue
        t_nospace = t.replace(" ", "").replace("\u3000", "")
        if t_nospace in non_outline_front_titles:
            continue

        target_level = None

        if t_nospace in special_h1_set:
            target_level = 1
        elif _match_chapter_heading(t, cfg, text_first=text_first):
            target_level = 1
        elif appendix_re.match(t):
            target_level = 1
        elif _looks_like_auto_heading_text(t, 2):
            if h4_re.match(t):
                target_level = 4
            elif h3_re.match(t):
                target_level = 3
            elif h2_re.match(t):
                target_level = 2

        if target_level is None:
            continue
        if target_level != 1 and not _looks_like_auto_heading_text(t, target_level):
            continue
        if target_level == 1 and t_nospace not in special_h1_set and not _looks_like_auto_heading_text(t, 1):
            continue
        if not _layout_allows_heading_level(para, target_level):
            continue
        if target_level == level:
            _clear_paragraph_outline_level(para)
            continue

        style = get_heading_style(doc, target_level)
        if style is None:
            continue
        if preserve_look:
            _materialize_paragraph_look(para)
        para.style = style
        _clear_paragraph_outline_level(para)
        changes.append(f'  {style.name}: "{t}"')

    return changes
